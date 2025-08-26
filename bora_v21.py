import os
import json
import shutil
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkinter.ttk import Progressbar
import time
from datetime import datetime
from pathlib import Path
import winshell
import fitz  # PyMuPDF
import docx
from docx      import Document
from docx.enum.text import WD_COLOR_INDEX
import xlrd
from bs4 import BeautifulSoup
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import sys
import subprocess
from html import unescape  # Used to convert HTML entities
import tkinter.font as tkfont  # For font manipulation in the text widget
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime as _dt
import winsound
import re

# Set up paths
CACHE_DIR = Path(__file__).parent / "cache"
SHORTCUT_DIR = Path(__file__).parent
CACHE_FILE = CACHE_DIR / "cache.json"
DASHBOARD_FILE = CACHE_DIR / "dashboard.json"

SUPPORTED_EXTENSIONS = [".txt", ".pdf", ".docx", ".xls", ".html"]
CACHE_DIR.mkdir(exist_ok=True)

# Define default width for the cached directories sidebar
DEFAULT_SIDEBAR_WIDTH = 150

#define words to ignore (stop-words)
COMMON_WORDS = {"a", "an", "the", "and", "or", "of", "in", "on", "for", "to", "with", "at", "by", "from", "up", "out", "over", "is", "isn't",
"has", "billion", "part", "parts",
"1", "2", "3", "4", "5", "6", "7", "8", "9", "0",
"-", "$", "&", "–", "'", "_",
"first", "two",
"I", "you", "we", "them",
"how", "what", "why", "where", "would", "here",
"solve", "as", "big", "small", "lull", "agents", "agent", "after", "before",
"de", "para", "o", "e", "que", "em", "no", "na", "agora", "os", "um", "é", "mais",
"ein",
","}

# right below your imports/constants
DEFAULT_AMPLIFIERS = {
    "words": ["pensascaramouche", "fuchspensamentos", "fuchs", "mentescrita", "§"]}

# Define a creation flag for subprocess calls (Windows)
creationflags = subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0

def open_file(path):
    """Opens a file or folder using the OS default application with no extra CMD window."""
    try:
        if os.name == "nt":
            os.startfile(path)
        elif sys.platform == "darwin":
            subprocess.call(("open", path), creationflags=creationflags)
        else:
            subprocess.call(("xdg-open", path), creationflags=creationflags)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to open: {e}")

def check_for_color_highlights(path):
    ext = os.path.splitext(path)[1].lower()
    if ext != ".pdf":
        return "N/A"
    try:
        doc = fitz.open(path)
        for page in doc:
            annots = page.annots()
            if annots is None:
                continue
            for annot in annots:
                if annot.type[0] == 8:
                    return "yes"
        return "no"
    except Exception as e:
        return f"error: {e}"

def retrieve_color_highlights(path):
    ext = os.path.splitext(path)[1].lower()
    if ext != ".pdf":
        return []
    highlights = []
    try:
        doc = fitz.open(path)
        for page in doc:
            annots = page.annots()
            if annots is None:
                continue
            for annot in annots:
                if annot.type[0] == 8:  # Highlight annotation
                    highlight_text = annot.info.get("content", "").strip()
                    if not highlight_text:
                        r = annot.rect
                        highlight_text = page.get_text("text", clip=r).strip()
                    colors = annot.colors
                    if "fill" in colors and colors["fill"]:
                        color_tuple = colors["fill"]
                    elif "stroke" in colors and colors["stroke"]:
                        color_tuple = colors["stroke"]
                    else:
                        color_tuple = (0, 0, 0)
                    try:
                        r_val, g_val, b_val = color_tuple
                        if r_val <= 1:
                            r_int = int(r_val * 255)
                            g_int = int(g_val * 255)
                            b_int = int(b_val * 255)
                        else:
                            r_int, g_int, b_int = int(r_val), int(g_val), int(b_val)
                        color_hex = f"#{r_int:02x}{g_int:02x}{b_int:02x}"
                    except Exception:
                        color_hex = "#000000"
                    highlights.append((color_hex, highlight_text))
        return highlights
    except Exception as e:
        print(f"[ERROR] retrieve_color_highlights: {e}")
        return []

def extract_year_key(path):
    """
    Return a tuple ((year, month, day), lowercase_name) for sorting,
    pushing non-dates to the end.
    """
    name = os.path.basename(path)
    bare = os.path.splitext(name)[0]
    tok  = bare.split()[0]
    parsed = _parse_token_date(tok)
    if parsed:
        y, mo, da = parsed
        date_key = (y, mo, da)
    else:
        date_key = (float('inf'), 0, 0)
    return (date_key, name.lower())

def compute_age(publish_date):
    now = _dt.now()
    years = now.year - publish_date.year
    months = now.month - publish_date.month
    days = now.day - publish_date.day
    if days < 0:
        months -= 1
        prev_month = now.month - 1 or 12
        prev_year = now.year if now.month > 1 else now.year - 1
        days_in_prev = (_dt(prev_year, now.month, 1) - _dt(prev_year, prev_month, 1)).days
        days += days_in_prev
    if months < 0:
        years -= 1
        months += 12
    return years, months, days		
		
def parse_css_style(style_str):
    """Parses a CSS style string and returns a dictionary of property-value pairs."""
    style_dict = {}
    if style_str:
        properties = style_str.split(";")
        for prop in properties:
            if ":" in prop:
                key, value = prop.split(":", 1)
                style_dict[key.strip().lower()] = value.strip()
    return style_dict

def retrieve_html_highlights(path):
    """
    For an HTML file, returns a list of tuples (color_hex, highlighted_text, style_dict)
    by extracting <span> elements whose style attribute contains 'background-color:'.
    """
    highlights = []
    try:
        content = Path(path).read_text(errors="ignore")
        soup = BeautifulSoup(content, "html.parser")
        spans = soup.find_all("span", style=lambda s: s and "background-color:" in s.lower())
        for span in spans:
            style = span.get("style", "")
            style_dict = parse_css_style(style)
            try:
                color_part = style.lower().split("background-color:")[1]
                color_hex = color_part.split(";")[0].strip()
            except Exception:
                color_hex = "#ffff99"
            text = unescape(span.get_text(strip=True))
            if text:
                highlights.append((color_hex, text, style_dict))
        return highlights
    except Exception as e:
        print(f"[ERROR] retrieve_html_highlights: {e}")
        return []

HIGHLIGHT_TO_HEX = {
    WD_COLOR_INDEX.YELLOW     : "#FFFF00",
    WD_COLOR_INDEX.BRIGHT_GREEN: "#00FF00",
    WD_COLOR_INDEX.TURQUOISE  : "#40E0D0",
    WD_COLOR_INDEX.PINK       : "#FFC0CB",
    WD_COLOR_INDEX.BLUE       : "#0000FF",
    WD_COLOR_INDEX.RED        : "#FF0000",
    WD_COLOR_INDEX.DARK_BLUE  : "#00008B",
    WD_COLOR_INDEX.TEAL       : "#008080",
    WD_COLOR_INDEX.GREEN      : "#008000",
    WD_COLOR_INDEX.VIOLET     : "#EE82EE",
    WD_COLOR_INDEX.DARK_RED   : "#8B0000",
    WD_COLOR_INDEX.DARK_YELLOW: "#CCCC00",
    WD_COLOR_INDEX.GRAY_50    : "#808080",
    WD_COLOR_INDEX.GRAY_25    : "#C0C0C0",
    WD_COLOR_INDEX.BLACK      : "#000000",
    WD_COLOR_INDEX.WHITE      : "#FFFFFF",
}

def retrieve_docx_highlights(path):
    """
    Extracts highlighted text from a DOCX file, including highlighted hyperlink text,
    and merges adjacent runs with the same highlight color into one block.
    Returns a list of (hex_color, text, style_dict).
    """
    highlights = []
    try:
        doc = Document(path)

        def process_runs(runs):
            merged_color = None
            merged_text = []
            for run in runs:
                color_enum = run.font.highlight_color
                text = run.text.strip()
                if color_enum is None or not text:
                    # flush if we're in a block
                    if merged_color is not None and merged_text:
                        hex_color = HIGHLIGHT_TO_HEX.get(merged_color, "#FFFF00")
                        style_dict = {
                            "highlight_index": merged_color,
                            "highlight_name": getattr(merged_color, "name", str(merged_color)),
                        }
                        highlights.append((hex_color, " ".join(merged_text), style_dict))
                        merged_color, merged_text = None, []
                    continue

                # If new block or color changes
                if merged_color is None or color_enum != merged_color:
                    if merged_color is not None and merged_text:
                        hex_color = HIGHLIGHT_TO_HEX.get(merged_color, "#FFFF00")
                        style_dict = {
                            "highlight_index": merged_color,
                            "highlight_name": getattr(merged_color, "name", str(merged_color)),
                        }
                        highlights.append((hex_color, " ".join(merged_text), style_dict))
                    merged_color = color_enum
                    merged_text = [text]
                else:
                    merged_text.append(text)

            # flush last
            if merged_color is not None and merged_text:
                hex_color = HIGHLIGHT_TO_HEX.get(merged_color, "#FFFF00")
                style_dict = {
                    "highlight_index": merged_color,
                    "highlight_name": getattr(merged_color, "name", str(merged_color)),
                }
                highlights.append((hex_color, " ".join(merged_text), style_dict))

        # Go through paragraphs
        for para in doc.paragraphs:
            process_runs(para.runs)

        # And through table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        process_runs(p.runs)

        return highlights

    except Exception as e:
        print(f"[ERROR] retrieve_docx_highlights: {e}")
        return []
		
def retrieve_pdf_highlights(path):
    """
    Extracts highlighted text and their colors from a PDF file.
    Returns a list of (hex_color, text, style_dict).
    """
    highlights = []

    try:
        doc = fitz.open(path)

        for page in doc:
            for annot in page.annots(types=[fitz.PDF_ANNOT_HIGHLIGHT]):
                quadpoints = annot.vertices
                text = ""
                if quadpoints:
                    words = page.get_text("words")  # list of words on the page
                    for i in range(0, len(quadpoints), 4):
                        rect = fitz.Quad(quadpoints[i:i+4]).rect
                        for w in words:
                            word_rect = fitz.Rect(w[:4])
                            if rect.intersects(word_rect):
                                text += w[4] + " "
                    text = text.strip()
                
                color = annot.colors["stroke"] if annot.colors else (1, 1, 0)  # RGB default yellow
                hex_color = '#%02x%02x%02x' % tuple(int(c * 255) for c in color)

                if text:
                    style_dict = {
                        "annot_type": "highlight",
                        "color_rgb": color,
                    }
                    highlights.append((hex_color, text, style_dict))

        return highlights

    except Exception as e:
        print(f"[ERROR] retrieve_pdf_highlights: {e}")
        return []		
		
def build_font_from_style(style_dict, base_font=("Courier", 10)):
    """Construct a tkfont.Font from CSS style properties."""
    family = style_dict.get("font-family", base_font[0])
    size_str = style_dict.get("font-size", None)
    try:
        size = int(size_str.replace("px", "").strip()) if size_str else base_font[1]
    except Exception:
        size = base_font[1]
    weight = style_dict.get("font-weight", "").lower()
    font_weight = "bold" if "bold" in weight else "normal"
    slant_val = style_dict.get("font-style", "").lower()
    font_slant = "italic" if "italic" in slant_val else "roman"
    return tkfont.Font(family=family, size=size, weight=font_weight, slant=font_slant)
	
def _parse_token_date(token):
    """
    Accepts YYYYMMDD, YYYY, -YYYYMMDD, or -YYYY and returns (year, month, day),
    or None if it doesn’t match.
    """
    # full 8-digit ±YYYYMMDD
    m = re.fullmatch(r'(-?\d{8})', token)
    if m:
        s = m.group(1)
        sign = -1 if s.startswith('-') else 1
        digits = s.lstrip('-')
        return (
            int(digits[0:4]) * sign,
            int(digits[4:6]),
            int(digits[6:8])
        )
    # just 4-digit ±YYYY
    m = re.fullmatch(r'(-?\d{4})', token)
    if m:
        s = m.group(1)
        sign = -1 if s.startswith('-') else 1
        return (int(s.lstrip('-')) * sign, 1, 1)
    return None

class DocumentRecommenderApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Document Recommender")
        self.root.geometry("800x700")
        self.root.rowconfigure(0, weight=1)
        self.root.columnconfigure(0, weight=1)
        self.sidebar_width = DEFAULT_SIDEBAR_WIDTH
		
        # ✅ add a first-run default so create_widgets() never crashes
        self.doc_info_font_size = 10

        # Mapping of directory -> timestamp
        self.directories = {}   # Persisted in dashboard.json
        # Mapping of directory -> cached file data
        self.cache = {}
        # Mapping of directory -> frame widget
        self.directory_frames = {}
        # Dictionary for toggle state (True=include, False=skip)
        self.directory_active_status = {}
        # Dictionary for reference to toggle button per directory.
        self.toggle_buttons = {}
        # Add this line alongside your other dicts:
        self.toggle_frames = {}   # holds the Frame for each directory’s toggle+erase buttons
        # Flag to remember if default directory has been used for file selection already.
        self.default_directory_used = False

        self.silencers = {"words": [], "factor": 1.0}
        self.default_directory = ""
        self.saved_files = set()
        self.load_dashboard()  # Load persisted dashboard & directories		
        # **wipe and load** the default amplifiers
        #self.amplifiers = {"words": DEFAULT_AMPLIFIERS["words"].copy(), "factor": 2.0}		
        #self.amplifiers = {"words": [], "factor": 1.0}

        self.current_document = None
        # This will store the computed recommendation tuples as (score, path)
        self.recommended_docs = []
        # BooleanVar for toggling display order
        self.sort_by_filename = tk.BooleanVar(value=False)
        # Keep track of the currently active file (path) and its label.
        self.current_active_file = None
        self.current_active_label = None

        self._rec_font_normal = ("Helvetica", 10, "normal")
        self._rec_font_bold   = ("Helvetica", 10, "bold")

        self.setup_scrollable_frame()
        self.create_widgets()
        self.load_cache()  # Build directory frames from cache

    def setup_scrollable_frame(self):
        # 1) Create a scrollable container for the _right_ pane only
        self.canvas = tk.Canvas(self.root)
        self.v_scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=self.v_scrollbar.set)
        self.v_scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)

        # 2) Host frame inside the canvas
        self.scroll_container = tk.Frame(self.canvas)

        # --- ADD these three lines to split content into sidebar + main area ---
        self.paned = ttk.Panedwindow(self.scroll_container, orient="horizontal")
        self.paned.pack(fill="both", expand=True)

        # use loaded width, then re-configure it before packing children
        self.sidebar = ttk.Frame(self.paned, width=self.sidebar_width, relief="sunken")
        self.sidebar.configure(width=self.sidebar_width)
        self.paned.add(self.sidebar, weight=0)

        self.right_pane = ttk.Frame(self.paned)
        self.paned.add(self.right_pane, weight=1)
        self.paned.sashpos(0, self.sidebar_width)
        self.root.after_idle(lambda: self.paned.sashpos(0, self.sidebar_width))

        # Whenever the user drags the sash, update our stored width:
        def on_sash_drag(event=None):
            # index 0 is the only sash in a two‑pane Panedwindow
            new_w = self.paned.sashpos(0)
            self.sidebar_width = new_w
            self.save_dashboard()

        # bind left‑mouse drag on the sash handle
        self.paned.bind("<B1-Motion>",     on_sash_drag)
        self.paned.bind("<ButtonRelease-1>", on_sash_drag)
        # --- END ADD ---

        # 3) Put our scroll_container into the canvas _once_, and never grid it
        self.canvas_window = self.canvas.create_window((0, 0), window=self.scroll_container, anchor="nw")
        self.scroll_container.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.bind("<Configure>", self.on_canvas_configure)

    def on_font_size_change(self):
        new_size = self.font_size_var.get()
        self.doc_info_font_size = new_size
        # Reconfigure the Text widget:
        current_font = tkfont.Font(font=self.output_text.cget("font"))
        current_font.configure(size=new_size)
        self.output_text.configure(font=current_font)
        # Persist:
        self.save_dashboard()

    def on_canvas_configure(self, event):
        # Expand our embedded window to fill both width and height:
        self.canvas.itemconfig(
            self.canvas_window,
            width=event.width,
            height=event.height
        )

    def create_widgets(self):
        # ——— Sidebar: Cached Directories ———
        tk.Label(self.sidebar, text="Cached Directories:", font=("Helvetica", 11, "bold")) \
            .pack(anchor="w", padx=5, pady=(5,0))

        # just above the Reload All button: a single toggle button instead of a Combobox
        self.view_mode = tk.StringVar(value="Cached tree")
        # create the button, initial color for “Cached tree”
        self.view_mode_button = tk.Button(
            self.sidebar,
            text=self.view_mode.get(),
            bg="#ccffcc",   # light green for Cached tree
            width=12,
            command=self.toggle_view_mode
        )
        self.view_mode_button.pack(anchor="w", padx=5, pady=(5,0))

        # Global “Reload All” button
        tk.Button(
            self.sidebar,
            text="Reload All Directories",
            command=self.reload_all_directories,
            bg="#ffd700"
        ).pack(anchor="e", padx=5, pady=(0,5))

        # ——— Search Bar ———
        search_frame = tk.Frame(self.sidebar)
        search_frame.pack(fill="x", padx=5, pady=(0,5))
        tk.Label(search_frame, text="Search (ayo):", font=("Helvetica", 10)).pack(side="left")
        self.search_var = tk.StringVar()
        search_entry = tk.Entry(search_frame, textvariable=self.search_var)
        search_entry.pack(side="left", fill="x", expand=True, padx=(5,0))
        # only search on Enter:
        search_entry.bind("<Return>", self.on_search_execute)
        # new “Go” button:
        tk.Button(search_frame, text="Go", command=self.on_search_execute).pack(side="right", padx=(5,0))

        # Frame to contain tree + its own scrollbar
        self.sidebar_tree_frame = ttk.Frame(self.sidebar)
        self.sidebar_tree_frame.pack(fill="both", expand=True, padx=5, pady=5)

        # 1) Make the scrollbar first
        self.tree_vscroll = ttk.Scrollbar(self.sidebar_tree_frame, orient="vertical")
        self.tree_vscroll.pack(side="right", fill="y")

        # 2) Make the Treeview after the scrollbar exists
        self.dir_tree = ttk.Treeview(
            self.sidebar_tree_frame,
            show="tree",
            yscrollcommand=self.tree_vscroll.set  # no lambda needed
        )
        self.dir_tree.pack(side="left", fill="both", expand=True)

        # 3) Wire scrollbar -> tree
        self.tree_vscroll.config(command=self.dir_tree.yview)

        # Make sure tags are configured
        self.dir_tree.tag_configure("active",   background="#ffffcc")
        self.dir_tree.tag_configure("inactive", background="#ffcccc")

        # Bind tree events
        self.dir_tree.bind("<Button-3>", self._on_tree_right_click)
        self.dir_tree.bind("<<TreeviewOpen>>",   self._on_tree_expand)
        self.dir_tree.bind("<ButtonRelease-1>",  self._on_tree_single_click)
        self.dir_tree.bind("<Double-1>",         self._on_tree_double_click)

        # ——— Now insert toggle buttons + erase ———
        tk.Label(self.sidebar, text="Toggle Directories:", font=("Helvetica", 10, "bold")) \
            .pack(anchor="w", padx=5, pady=(10,0))

        self.toggle_btns_frame = ttk.Frame(self.sidebar)
        self.toggle_btns_frame.pack(fill="x", padx=5, pady=(0,5))

        self.refresh_toggle_buttons()

        # Create the context‐menu once
        self.dir_menu = tk.Menu(self.sidebar, tearoff=0)
        self.dir_menu.add_command(label="Save On/Off", command=self._toggle_save_selected_file)
        self.dir_menu.add_command(label="RUN", command=self._run_selected_file)
        self.dir_menu.add_command(label="Folder", command=self._open_selected_file)
        self.dir_menu.add_command(label="Zoom", command=self._zoom_selected_dirs)
        #self.dir_menu.add_separator()
        self.dir_menu.add_command(label="Copy path", command=self._copy_path_to_clipboard)

        # Bind tree events once
        self.dir_tree.bind("<Button-3>", self._on_tree_right_click)
        self.dir_tree.bind("<<TreeviewOpen>>", self._on_tree_expand)
        self.dir_tree.bind("<ButtonRelease-1>", self._on_tree_single_click)
        self.dir_tree.bind("<Double-1>", self._on_tree_double_click)


        # --- END ADD ---
	
     	# Top Panel: Default Directory and Control Buttons in one compact row.
        top_panel = tk.Frame(self.right_pane, bd=2, relief="raised", padx=5, pady=5)
        top_panel.pack(fill="x", pady=(5,10))
        tk.Label(top_panel, text="Default Dir:", font=("Helvetica", 10, "bold")).pack(side="left", padx=3)
        self.default_dir_entry = tk.Entry(top_panel, width=40)
        self.default_dir_entry.pack(side="left", padx=3)
        self.default_dir_entry.insert(0, self.default_directory)
        tk.Button(top_panel, text="Save", command=self.save_default_directory, bg="#ccffcc").pack(side="left", padx=3)
        tk.Button(top_panel, text="Add Dir", command=self.add_directory, bg="#cce5ff").pack(side="left", padx=3)
        tk.Button(top_panel, text="Select File", command=self.select_current_document, bg="#ffebcc").pack(side="left", padx=3)
        tk.Button(top_panel, text="Rerun", command=self.run_dashboard, bg="#ffcccc").pack(side="left", padx=3)
        self.status_label = tk.Label(top_panel, text="", font=("Helvetica", 9, "italic"))
        self.status_label.pack(side="left", padx=10)
		
        self.toggle_rec_btn = tk.Button(
            top_panel,
            text="Hide",
            command=self.toggle_recommendations,
            bg="#d3d3d3"
        )
        self.toggle_rec_btn.pack(side="right", padx=3)
        
        # Middle Section: Document Info and Recommendation Panels.
        middle_frame = tk.Frame(self.right_pane)
        middle_frame.pack(fill="both", expand=True, pady=10)
        middle_frame.grid_columnconfigure(0, weight=1, uniform="halves")
        middle_frame.grid_columnconfigure(1, weight=1, uniform="halves")
        middle_frame.grid_rowconfigure(0, weight=1)
        
        self.about_frame = tk.Frame(middle_frame, bd=1, relief="groove")
        self.about_frame.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)

        # ── Header: “Document Info:” + current document title ──
        header_frame = tk.Frame(self.about_frame)
        header_frame.pack(fill="x", padx=5, pady=5)
        tk.Label(header_frame, text="Document Info:", font=("Helvetica", 12, "bold")).pack(side="left")
        # this label will show the current file name
        self.doc_info_title_label = tk.Label(header_frame, text="No file selected", font=("Helvetica", 10, "bold"))
        self.doc_info_title_label.pack(side="left", padx=(5,0))
		
        # Right after self.doc_info_title_label.pack(...)
        tk.Label(header_frame, text="Font size:").pack(side="left", padx=(10,0))
        self.font_size_var = tk.IntVar(value=self.doc_info_font_size)
        font_spin = tk.Spinbox(
            header_frame,
            from_=6, to=32,
            textvariable=self.font_size_var,
            width=3,
            command=self.on_font_size_change
        )
        font_spin.pack(side="left")

        # container to hold text widget + its scrollbar
        text_frame = tk.Frame(self.about_frame)
        text_frame.pack(fill="both", expand=True, padx=5, pady=5)

        # 1) scrollbar on the left
        scroll_bar = ttk.Scrollbar(text_frame, orient="vertical")
        scroll_bar.pack(side="left", fill="y")

        # 2) the Text widget bound to that scrollbar
        self.output_text = tk.Text(
            text_frame,
            wrap="word",
            font=("Verdana", self.doc_info_font_size),
            yscrollcommand=scroll_bar.set
        )
        self.output_text.pack(side="left", fill="both", expand=True)

        # hook them together
        scroll_bar.config(command=self.output_text.yview)

        # 3) ensure scrollbar is on top of the stacking order
        scroll_bar.lift()

        self.output_text.config(state="disabled")
        
        self.rec_frame = tk.Frame(middle_frame, bd=1, relief="groove")
        self.rec_frame.grid(row=0, column=1, sticky="nsew", padx=5, pady=5)
        tk.Label(self.rec_frame, text="Recommendations:", font=("Helvetica", 12, "bold")).pack(anchor="w", padx=5, pady=5)
        tk.Checkbutton(self.rec_frame, text="Sort by Filename", variable=self.sort_by_filename, 
                       command=self.refresh_recommendation_buttons).pack(anchor="w", padx=5)
        self.shortcut_buttons_frame = tk.Frame(self.rec_frame)
        self.shortcut_buttons_frame.pack(fill="both", expand=True, padx=5, pady=5)

        # Compact Dashboard Section: Amplifiers & Silencers in a single row.
        self.dashboard_frame = tk.Frame(self.right_pane, bd=2, relief="sunken", padx=5, pady=5)
        self.dashboard_frame.pack(fill="x", pady=(5,0), padx=5)
        tk.Label(self.dashboard_frame, text="Dashboard", font=("Helvetica", 11, "bold")).grid(row=0, column=0, columnspan=2, sticky="w")

        # Amplifiers Panel (colored light blue)
        amp_frame = tk.Frame(self.dashboard_frame, bg="#e0f7ff", padx=3, pady=3)
        amp_frame.grid(row=1, column=0, sticky="nsew", padx=3, pady=3)
        tk.Label(amp_frame, text="Amp", bg="#e0f7ff", font=("Helvetica", 9, "bold")).grid(row=0, column=0, sticky="w", padx=2)
        tk.Label(amp_frame, text="F:", bg="#e0f7ff").grid(row=0, column=1, sticky="w")
        self.amp_factor_entry = tk.Entry(amp_frame, width=4)
        self.amp_factor_entry.insert(0, str(self.amplifiers["factor"]))
        self.amp_factor_entry.grid(row=0, column=2, sticky="w", padx=2)
        tk.Label(amp_frame, text="Word:", bg="#e0f7ff").grid(row=0, column=3, sticky="w")
        self.amp_word_entry = tk.Entry(amp_frame, width=10)
        self.amp_word_entry.grid(row=0, column=4, sticky="w", padx=2)
        tk.Button(amp_frame, text="+", command=self.add_amplifier_word, width=2, bg="#b3e6ff").grid(row=0, column=5, sticky="w", padx=2)
        self.amp_list_frame = tk.Frame(amp_frame, bg="#e0f7ff")
        self.amp_list_frame.grid(row=1, column=0, columnspan=6, sticky="w", padx=2, pady=2)

        # Silencers Panel (colored light coral)
        sil_frame = tk.Frame(self.dashboard_frame, bg="#ffe0e0", padx=3, pady=3)
        sil_frame.grid(row=1, column=1, sticky="nsew", padx=3, pady=3)
        tk.Label(sil_frame, text="Sil", bg="#ffe0e0", font=("Helvetica", 9, "bold")).grid(row=0, column=0, sticky="w", padx=2)
        tk.Label(sil_frame, text="F:", bg="#ffe0e0").grid(row=0, column=1, sticky="w")
        self.sil_factor_entry = tk.Entry(sil_frame, width=4)
        self.sil_factor_entry.insert(0, str(self.silencers["factor"]))
        self.sil_factor_entry.grid(row=0, column=2, sticky="w", padx=2)
        tk.Label(sil_frame, text="Word:", bg="#ffe0e0").grid(row=0, column=3, sticky="w")
        self.sil_word_entry = tk.Entry(sil_frame, width=10)
        self.sil_word_entry.grid(row=0, column=4, sticky="w", padx=2)
        tk.Button(sil_frame, text="+", command=self.add_silencer_word, width=2, bg="#ffb3b3").grid(row=0, column=5, sticky="w", padx=2)
        self.sil_list_frame = tk.Frame(sil_frame, bg="#ffe0e0")
        self.sil_list_frame.grid(row=1, column=0, columnspan=6, sticky="w", padx=2, pady=2)

    def _open_selected_file(self):
        selected_item = self.dir_tree.focus()
        if selected_item:
            # Get the full path from the selected item's values
            # Assumes the full path is stored as the first value in the Treeview item
            file_path = self.dir_tree.item(selected_item, "values")[0]

            if os.path.isfile(file_path):
                folder_path = os.path.dirname(file_path)
            elif os.path.isdir(file_path): # If the selected item is already a directory
                folder_path = file_path
            else:
                messagebox.showerror("Error", "Selected item is neither a file nor a valid directory.")
                return

            if folder_path:
                try:
                    # Open the folder using the default file explorer
                    os.startfile(folder_path)
                except Exception as e:
                    messagebox.showerror("Error", f"Could not open folder:\n{e}")

    def toggle_recommendations(self):
        """Show/hide recs and let the about pane span the full width when hidden."""
        if self.rec_frame.winfo_ismapped():
            # hide the rec_frame
            self.rec_frame.grid_remove()
            # make about_frame take both columns
            self.about_frame.grid_configure(columnspan=2)
            self.toggle_rec_btn.config(text="Show")
        else:
            # restore rec_frame
            self.rec_frame.grid()
            # put about_frame back to single column
            self.about_frame.grid_configure(columnspan=1)
            self.toggle_rec_btn.config(text="Hide")

    def ensure_recommendations_visible(self):
        # If hidden, re-show the rec_frame and reset the toggle button:
        if not self.rec_frame.winfo_ismapped():
            self.rec_frame.grid()
            self.about_frame.grid_configure(columnspan=1)
            self.toggle_rec_btn.config(text="Hide")

    def _zoom_selected_dirs(self):
        """
        Show a flat, alphabetically sorted list of all files
        in the currently active (toggled-on) directories,
        and center the right-clicked item in view.
        """
        # Remember the path of the right-clicked node
        focused = self.dir_tree.focus()
        target_path = self._node_path(focused) if focused else None
		
        # parse the base date from the zoomed file’s name
        base_y, base_m, base_d = (None, None, None)
        if target_path:
            base_name = os.path.splitext(os.path.basename(target_path))[0]
            tok = base_name.split()[0]
            parsed = _parse_token_date(tok)
            if parsed:
                base_y, base_m, base_d = parsed

        # 1) Collect files from active directories
        all_files = []
        # Regex to check if a filename starts with at least 4 digits
        # OR starts with a minus sign followed by exactly 4 digits
        filename_pattern = re.compile(r"^(?:\d{4,}|-\d{4})") # MODIFIED LINE
        for directory, is_active in self.directory_active_status.items():
            if not is_active:
                continue
            for root, _, filenames in os.walk(directory):
                for fname in filenames:
                    # Apply the filter here: only add if filename matches the pattern
                    if filename_pattern.match(os.path.basename(fname)): # Check against the new pattern
                        file_path = os.path.join(root, fname)
                        all_files.append(file_path)
						
        # 2) Sort by numeric year (BC dates negative) and then filename
        all_files.sort(key=extract_year_key)

        # 3) Clear existing tree
        self.dir_tree.delete(*self.dir_tree.get_children())

        path_to_iid = {}

        # 4) Insert flat file list, but now prefix with (years,months)
        for path in all_files:
            name = os.path.basename(path)
            # build the prefix from (years,months) relative to zoom base
            prefix = ""
            if base_y is not None:
                base = (base_y, base_m, base_d)

                # strip extension and grab token
                bare = os.path.splitext(name)[0]
                tok = bare.split()[0]
                dvals = _parse_token_date(tok)

                if dvals:
                    y, mo, da = dvals
                    # raw differences
                    yrs  = y  - base[0]
                    mos  = mo - base[1]
                    days = da  - base[2]

                    # if day < 0, we haven’t hit that month anniversary
                    if days < 0:
                        mos -= 1
                    # if mos < 0, borrow from years
                    if mos < 0:
                        yrs -= 1
                        mos += 12

                    prefix = f"{yrs},{mos} "


            iid = self.dir_tree.insert(
                "", "end",
                text=prefix + name,
                values=(path,)
            )
            path_to_iid[path] = iid

        # 5) If the right-clicked item is among these files, select and scroll it to top
        if target_path in path_to_iid:
            iid = path_to_iid[target_path]

            # Highlight selection
            self.dir_tree.selection_set(iid)
            self.dir_tree.focus(iid)

            # Scroll the item into view, attempting to place it at the top
            self.dir_tree.see(iid)

            # Optional: You might want to adjust scrolling slightly if 'see' doesn't put it perfectly at the top
            # This is more of a fine-tuning if 'see' alone isn't sufficient for your exact visual preference.
            # You could try scrolling up by a small fraction if the item is still not at the very top.
            # For example, to scroll up by one "line" if needed:
            # self.dir_tree.yview_scroll(-1, "units")
            # However, for placing it at the top, 'see(iid)' is generally the best approach.

        # Update view mode to reflect "Zoomed" state
        self.view_mode.set("Zoomed")
        self.view_mode_button.config(text="Zoomed", bg="#add8e6") # Light blue for zoomed mode

    def on_search_execute(self, event=None):
        query = self.search_var.get().strip()
        if not query:
            self.on_view_mode_change()
            return

        # --- detect & strip "ayo" trigger ---
        year_sort = False
        q_lower = query.lower()
        if q_lower.endswith(" ayo"):
            year_sort = True
            query = query[:-len(" ayo")].rstrip()
        elif q_lower == "ayo":
            year_sort = True
            query = ""
        # --------------------------------------

        # Parse into phrases & words
        search_parts = []
        for match in re.finditer(r'"([^"]*)"|(\S+)', query):
            if match.group(1):
                search_parts.append({"type": "phrase", "value": match.group(1)})
            else:
                search_parts.append({"type": "word", "value": match.group(2)})

        if not search_parts and not year_sort:
            self.on_view_mode_change()
            return

        matches_with_counts = []
        for directory, data in self.cache.items():
            if not self.directory_active_status.get(directory, True):
                continue

            for path, text in zip(data["files"], data["texts"]):
                filename = os.path.basename(path)
                total_occurrences = 0
                all_ok = True
                matched_in_filename = False

                for part in search_parts:
                    val = part["value"]
                    if part["type"] == "phrase":
                        pat = re.compile(re.escape(val), re.IGNORECASE)
                    else:
                        pat = re.compile(rf'\b{re.escape(val)}\b', re.IGNORECASE)

                    text_occ = len(pat.findall(text))
                    if text_occ > 0:
                        total_occurrences += text_occ
                    elif pat.search(filename):
                        matched_in_filename = True
                    else:
                        all_ok = False
                        break

                if all_ok:
                    # Count filename match as single occurrence if no content matches
                    if total_occurrences == 0 and matched_in_filename:
                        total_occurrences = 1
                    matches_with_counts.append((path, total_occurrences))

        # Sort: alphabetical if ayo, else by count desc
        if year_sort:
            matches_with_counts.sort(key=lambda x: extract_year_key(x[0]))
        else:
            matches_with_counts.sort(
                key=lambda x: (-x[1], os.path.basename(x[0]).lower())
            )

        # Rebuild tree
        self.dir_tree.delete(*self.dir_tree.get_children())
        for path, count in matches_with_counts:
            display = f"{os.path.basename(path)} ({count})"
            self.dir_tree.insert("", "end", text=display, values=(path,))

    def _run_selected_file(self):
        """Called by the right-click RUN menu item."""
        # 1) Which row?
        sel = self.dir_tree.selection()
        if not sel:
            return
        iid = sel[0]

        # 2) Get the real filesystem path
        path = self._node_path(iid)
        if not os.path.isfile(path):
            return

        # 3) “Run” it—same as selecting it via the GUI
        self.dig_document(path)

    def toggle_view_mode(self):
        # flip between modes
        new_mode = "Saved" if self.view_mode.get() == "Cached tree" else "Cached tree"
        self.view_mode.set(new_mode)

        # choose a distinct color for each mode
        if new_mode == "Cached tree":
            color = "#ccffcc"   # light green
        else:
            color = "#cce5ff"   # light blue

        # update the button’s appearance
        self.view_mode_button.config(text=new_mode, bg=color)

        # rebuild the tree
        self.on_view_mode_change()

    def on_view_mode_change(self, *_):
            mode = self.view_mode.get()
            # Always clear the tree when changing modes
            for iid in self.dir_tree.get_children():
                self.dir_tree.delete(iid)

            if mode == "Cached tree":
                # Rebuild the tree roots for each cached directory
                for directory in sorted(self.directories):
                    self.build_directory_tree(directory)
            elif mode == "Saved":
                # Build the tree specifically for saved files
                self.build_saved_files_tree()

    def build_saved_files_tree(self):
            # Clear existing items if any (already handled by on_view_mode_change, but good for self-contained logic)
            # for iid in self.dir_tree.get_children():
            #     self.dir_tree.delete(iid)

            # Sort saved files alphabetically by their base name for a flat list
            # We use os.path.basename for sorting key
            sorted_saved_files = sorted(list(self.saved_files), key=lambda x: os.path.basename(x).lower())

            for file_path in sorted_saved_files:
                normalized_file_path = str(Path(file_path).resolve())
                file_name = os.path.basename(normalized_file_path)

                # Insert the file directly as a top-level item
                # The 'values' attribute still stores the full path for later use (e.g., opening the file)
                self.dir_tree.insert("", "end", text=file_name, values=(normalized_file_path,))

    def _on_tree_single_click(self, event):
        iid = self.dir_tree.identify_row(event.y)
        if not iid:
            return
        path = self._node_path(iid)
        if os.path.isfile(path):
            # 1) Update the “current document” state and UI label
            self.current_document = path
            self.doc_info_title_label.config(text=os.path.basename(path))
            # 2) Extract its full text (same extractor used elsewhere)
            text = self.extract_text(path)
            # 3) Only load the info pane—do *not* kick off recommendations yet
            self.show_document_info(path, text)

    def _on_tree_double_click(self, event):
        iid = self.dir_tree.identify_row(event.y)
        if not iid:
            return
        path = self._node_path(iid)
        if os.path.isfile(path):
            open_file(path)  # not self.open_file

    def _on_sidebar_toggle(self, dir_path):
        # 1) Flip the status in your app logic
        self.toggle_directory(dir_path)

        # 2) Read back the new status
        new_status = self.directory_active_status.get(dir_path, False)

        # 3) Update only the button color
        btn = self.toggle_buttons.get(dir_path)
        if btn:
            btn.config(bg="green" if new_status else "red")

        # 4) Also recolor the tree node (if you want consistency there)
        for iid in self.dir_tree.get_children():
            # Normalize path for comparison
            if str(Path(self.dir_tree.item(iid, "values")[0]).resolve()) == str(Path(dir_path).resolve()):
                tag = "active" if new_status else "inactive"
                self.dir_tree.item(iid, tags=(tag,))
                break

    def _on_tree_expand(self, event):
        # original handler likely did something like:
        # sel = self.dir_tree.selection()[0]
        # now just:
        sel_list = self.dir_tree.selection()
        if not sel_list:
            return
        self._expand_and_populate(sel_list[0])

    def _on_tree_select(self, event):
        sel = self.dir_tree.selection()
        if not sel:
            return
        path = self._node_path(sel[0])
        if os.path.isfile(path):
            self.load_file(path)
				
    def _node_path(self, item):
        """Returns the full normalized path associated with a Treeview item."""
        item_values = self.dir_tree.item(item, "values")
        if item_values and item_values[0]:
            return str(Path(item_values[0]).resolve()) # Normalize path when retrieved
        return ""


    def reload_all_directories(self):
        """
        Reloads (rebuilds) the cache for every known directory.
        """
        # Disable UI updates during mass reload
        self.status_label.config(text="Reloading all directories...")
        for directory in list(self.directories.keys()):
            # Force rebuild cache for each directory
            self.build_cache(directory, force_rebuild=True)
        self.status_label.config(text="All directories reloaded.")
		
    def save_default_directory(self):
        # Normalize the default directory path before saving
        normalized_default_dir = str(Path(self.default_dir_entry.get().strip()).resolve()) if self.default_dir_entry.get().strip() else ""
        # Ensure trailing slash for root directories for consistency
        if Path(normalized_default_dir).parent == Path(normalized_default_dir) and not normalized_default_dir.endswith(os.sep):
            normalized_default_dir += os.sep
        self.default_directory = normalized_default_dir
        self.save_dashboard()
        messagebox.showinfo("Saved", "Default directory updated.")

    def add_amplifier_word(self):
        word = self.amp_word_entry.get().strip()
        if word and word not in self.amplifiers["words"]:
            self.amplifiers["words"].append(word)
        self.amp_word_entry.delete(0, tk.END)
        self.refresh_amp_list()
        self.save_dashboard()

    def refresh_amp_list(self):
        for widget in self.amp_list_frame.winfo_children():
            widget.destroy()
        num_cols = 3  # Display in a compact way.
        for idx, word in enumerate(self.amplifiers["words"]):
            col = idx % num_cols
            row = idx // num_cols
            frame = tk.Frame(self.amp_list_frame, bg="#e0f7ff")
            frame.grid(row=row, column=col, padx=2, pady=2, sticky="w")
            tk.Label(frame, text=word, bg="#e0f7ff").pack(side="left")
            tk.Button(frame, text="×", command=lambda i=idx: self.remove_amplifier_word(i), bg="#ffcccc", padx=1).pack(side="left", padx=2)

    def remove_amplifier_word(self, index):
        if 0 <= index < len(self.amplifiers["words"]):
            del self.amplifiers["words"][index]
            self.refresh_amp_list()
            self.save_dashboard()

    def add_silencer_word(self):
        word = self.sil_word_entry.get().strip()
        if word and word not in self.silencers["words"]:
            self.silencers["words"].append(word)
        self.sil_word_entry.delete(0, tk.END)
        self.refresh_sil_list()
        self.save_dashboard()

    def refresh_sil_list(self):
        for widget in self.sil_list_frame.winfo_children():
            widget.destroy()
        num_cols = 3
        for idx, word in enumerate(self.silencers["words"]):
            col = idx % num_cols
            row = idx // num_cols
            frame = tk.Frame(self.sil_list_frame, bg="#ffe0e0")
            frame.grid(row=row, column=col, padx=2, pady=2, sticky="w")
            tk.Label(frame, text=word, bg="#ffe0e0").pack(side="left")
            tk.Button(frame, text="×", command=lambda i=idx: self.remove_silencer_word(i), bg="#ffb3b3", padx=1).pack(side="left", padx=2)

    def remove_silencer_word(self, index):
        if 0 <= index < len(self.silencers["words"]):
            del self.silencers["words"][index]
            self.refresh_sil_list()
            self.save_dashboard()

    def run_dashboard(self):
        self.ensure_recommendations_visible()
        try:
            amp_factor = float(self.amp_factor_entry.get())
            sil_factor = float(self.sil_factor_entry.get())
        except ValueError:
            messagebox.showerror("Input Error", "Factors must be numbers.")
            return
        self.amplifiers["factor"] = amp_factor
        self.silencers["factor"] = sil_factor
        self.save_dashboard()
        if self.current_document:
            text = self.extract_text(self.current_document)
            self.show_document_info(self.current_document, text)
            self.clear_shortcut_buttons()
            self.recommend_similar_files(self.current_document)

    def save_dashboard(self):
        data = {
            "amplifiers": self.amplifiers,
            "silencers": self.silencers,
            "default_directory": self.default_directory,
            "directory_active_status": self.directory_active_status,
            "directories": self.directories,
            "saved_files": list(self.saved_files),
            "sidebar_width": self.sidebar_width,
            "doc_info_font_size": self.doc_info_font_size
        }
        with open(DASHBOARD_FILE, "w") as f:
            json.dump(data, f)

    def load_dashboard(self):
        if DASHBOARD_FILE.exists():
            try:
                with open(DASHBOARD_FILE, "r") as f:
                    data = json.load(f)

                self.amplifiers = data.get("amplifiers", {"words": [], "factor": 1.0})
                self.silencers = data.get("silencers", {"words": [], "factor": 1.0})
                self.saved_files = set(data.get("saved_files", []))

                # default directory (normalize + trailing slash for root)
                default_dir = data.get("default_directory", "")
                if default_dir:
                    norm = str(Path(default_dir).resolve())
                    if Path(norm).parent == Path(norm) and not norm.endswith(os.sep):
                        norm += os.sep
                    self.default_directory = norm
                else:
                    self.default_directory = ""

                # normalize path-keyed dicts
                normalized_active_status = {}
                for k, v in data.get("directory_active_status", {}).items():
                    nk = str(Path(k).resolve())
                    if Path(nk).parent == Path(nk) and not nk.endswith(os.sep):
                        nk += os.sep
                    normalized_active_status[nk] = v
                self.directory_active_status = normalized_active_status

                normalized_directories = {}
                for k, v in data.get("directories", {}).items():
                    nk = str(Path(k).resolve())
                    if Path(nk).parent == Path(nk) and not nk.endswith(os.sep):
                        nk += os.sep
                    normalized_directories[nk] = v
                self.directories = normalized_directories

                self.sidebar_width = data.get("sidebar_width", self.sidebar_width)
                self.doc_info_font_size = data.get("doc_info_font_size", 10)

            except Exception as e:
                print(f"[ERROR] Loading dashboard: {e}")
                # fall back to first-run defaults if the file is corrupt
                self._apply_first_run_defaults()
                self.save_dashboard()
        else:
            # first run: no file yet
            self._apply_first_run_defaults()
            self.save_dashboard()

    def _apply_first_run_defaults(self):
        self.amplifiers = {"words": DEFAULT_AMPLIFIERS["words"].copy(), "factor": 2.0}
        self.silencers = {"words": [], "factor": 1.0}
        self.saved_files = set()
        self.default_directory = ""
        self.directory_active_status = {}
        self.directories = {}
        self.sidebar_width = DEFAULT_SIDEBAR_WIDTH
        self.doc_info_font_size = 10

    def add_directory(self):
        directory = filedialog.askdirectory()
        if directory:
            # Normalize the path immediately upon selection
            normalized_directory = str(Path(directory).resolve()) 
            # Ensure trailing slash for root directories for consistency
            if Path(normalized_directory).parent == Path(normalized_directory) and not normalized_directory.endswith(os.sep):
                normalized_directory += os.sep

            if normalized_directory not in self.directories:
                self.directories[normalized_directory] = datetime.now().isoformat()
                self.directory_active_status[normalized_directory] = self.directory_active_status.get(normalized_directory, True)
                self.build_directory_tree(normalized_directory)
                self.build_cache(normalized_directory)
                self.save_dashboard()

    def build_directory_tree(self, directory):
        # directory here is already normalized.
        # Remove any existing node
        for iid in self.dir_tree.get_children():
            # Ensure comparison is also normalized
            if str(Path(self.dir_tree.item(iid, "values")[0]).resolve()) == str(Path(directory).resolve()):
                self.dir_tree.delete(iid)

        # Insert root node with normalized directory path
        # Use basename for display, full path in values
        display_text = os.path.basename(directory)
        if not display_text and Path(directory).is_dir() and Path(directory).drive: # It's a root drive like C:\
            display_text = directory # Use the full path for display if it's a root drive
        elif not display_text: # Fallback for non-drive roots (e.g., '/') or other edge cases
            display_text = directory

        iid = self.dir_tree.insert("", "end", text=display_text, values=(directory,)) 
        self.dir_tree.insert(iid, "end", text="…")  # dummy so ▸ appears

        tag = "active" if self.directory_active_status.get(directory, True) else "inactive"
        self.dir_tree.item(iid, tags=(tag,))
        # tag_configure should be done once, not repeatedly here.
        # self.dir_tree.tag_configure("active",   background="#ffffcc")
        # self.dir_tree.tag_configure("inactive", background="#ffcccc")

    def _on_tree_right_click(self, event):
        iid = self.dir_tree.identify_row(event.y)
        if not iid:
            return
        self.dir_tree.selection_set(iid)
        self.dir_menu.tk_popup(event.x_root, event.y_root)

    def _reload_selected_dir(self):
        d = self._node_path(self.dir_tree.selection()[0]) # Use normalized path
        self.build_cache(d, force_rebuild=True)

    def _erase_selected_dir(self):
        d = self._node_path(self.dir_tree.selection()[0]) # Use normalized path
        self.erase_directory(d)

    def _toggle_save_selected_file(self):
        iid = self.dir_tree.selection()[0]
        path = self._node_path(iid)
        if path in self.saved_files:
            self.saved_files.remove(path)
        else:
            self.saved_files.add(path)
        self.save_dashboard()
        # if we’re in “Saved” view, remove/add the node immediately:
        if self.view_mode.get() == "Saved":
            self.dir_tree.delete(iid)

    def _copy_path_to_clipboard(self):
        iid = self.dir_tree.selection()
        if not iid: return
        path = self._node_path(iid[0])
        # Use Tk’s clipboard API
        self.root.clipboard_clear()
        self.root.clipboard_append(path)
        messagebox.showinfo("Copied", f"Path copied to clipboard:\n{path}")

    def toggle_directory(self, directory):
        # directory here is already normalized.
        current = self.directory_active_status.get(directory, True)
        new_status = not current
        self.directory_active_status[directory] = new_status
        button = self.toggle_buttons.get(directory)
        if button:
            button.config(bg="green" if new_status else "red")
        frame = self.directory_frames.get(directory)
        if frame:
            if new_status:
                frame.config(bg="#ffffcc")
                for child in frame.winfo_children():
                    child.config(bg="#ffffcc")
            else:
                frame.config(bg="#ffcccc")
                for child in frame.winfo_children():
                    child.config(bg="#ffcccc")
        self.save_dashboard()
        print(f"[LOG] Directory '{directory}' set to {'active' if new_status else 'inactive'}.")

    def update_directory_frame(self, directory):
        frame = self.directory_frames.get(directory)
        if frame:
            for widget in frame.winfo_children():
                if isinstance(widget, tk.Label) and widget.cget("text").startswith("Loaded:"):
                    ts_str = self.directories.get(directory, None)
                    new_ts = self.format_timestamp(ts_str) if ts_str else "Never"
                    widget.config(text=f"Loaded: {new_ts}")
                    break

    def erase_directory(self, directory):
        # directory here is already normalized.
        if directory in self.directories:
            del self.directories[directory]
        if directory in self.cache:
            del self.cache[directory]
            self.save_cache()
        # You would also need to remove the corresponding Tkinter frames/buttons for the erased directory
        # The refresh_toggle_buttons should handle removal of old frames/buttons.
        # This part assumes refresh_toggle_buttons will be called after erasing.
        if directory in self.directory_active_status:
            del self.directory_active_status[directory]
        
        # Remove the directory from the Treeview
        for iid in self.dir_tree.get_children():
            if str(Path(self.dir_tree.item(iid, "values")[0]).resolve()) == str(Path(directory).resolve()):
                self.dir_tree.delete(iid)
                break
        
        self.refresh_toggle_buttons() # Rebuild toggle buttons after erase
        self.save_dashboard()
        print(f"[LOG] Erased directory: {directory}")

    def build_cache(self, directory, force_rebuild=False):
        """
        Incrementally rebuilds the cache for directory, only re-extracting files
        whose modification time has changed (or if force_rebuild=True).
        """
        # directory here is already normalized.
        # 1) Load old cache and detect if mtimes are available
        old = self.cache.get(directory, {})
        old_files = old.get("files", [])
        old_texts = old.get("texts", [])
        old_mtimes = old.get("mtimes", [])
        # If mtimes length mismatches, force full rebuild
        if len(old_mtimes) != len(old_files):
            force_rebuild = True

        # 2) Build lookup for reuse, ensuring keys are normalized paths
        old_index = {} if force_rebuild else {
            str(Path(fp).resolve()): (txt, mtime)
            for fp, txt, mtime in zip(old_files, old_texts, old_mtimes)
        }

        # 3) Discover all supported files recursively, normalizing paths
        filepaths = [
            str(Path(p).resolve()) for p in Path(directory).rglob("*")
            if p.suffix.lower() in SUPPORTED_EXTENSIONS
        ]
        total = len(filepaths)
        new_entries = []  # will hold tuples (path, mtime, text)

        # 4) Process each file: reuse or re-extract
        for idx, path in enumerate(filepaths, start=1):
            try:
                mtime = os.path.getmtime(path)
            except OSError:
                mtime = None

            if (not force_rebuild
                and path in old_index
                and old_index[path][1] == mtime):
                # reuse cached text
                text = old_index[path][0]
            else:
                # fresh extraction
                try:
                    text = self.extract_text(path)
                except Exception as e:
                    print(f"[ERROR] Failed to extract {path}: {e}")
                    text = ""

            new_entries.append((path, mtime, text))

            # Optional: update progress
            pct = int(idx / total * 100) if total else 100
            self.status_label.config(text=f"Processing {idx}/{total} ({pct}%)")
            self.root.update_idletasks()

        # 5) Unpack new entries into parallel lists
        if new_entries:
            files, mtimes, texts = zip(*new_entries)
        else:
            files, mtimes, texts = [], [], []

        # 6) Update cache and UI
        self.cache[directory] = { # directory key is already normalized here
            "files": list(files),
            "texts": list(texts),
            "mtimes": list(mtimes),
            "timestamp": datetime.now().isoformat()
        }
        self.directories[directory] = self.cache[directory]["timestamp"] # directory key is already normalized here
        self.save_cache()
        self.save_dashboard()
        self.update_directory_frame(directory)
        self.status_label.config(text="Cache rebuilt incrementally.")

    def save_cache(self):
        with open(CACHE_FILE, "w") as f:
            json.dump(self.cache, f)

    def load_cache(self):
        """
        Loads cached file data and ensures the GUI shows all directories
        present in cache.json, even if they weren’t explicitly added
        in the dashboard.
        """
        if CACHE_FILE.exists():
            try:
                # 1) Load raw cache
                with open(CACHE_FILE, "r") as f:
                    raw_cache = json.load(f)
                
                # 2) Normalize raw_cache keys to ensure consistency in self.cache
                normalized_cache = {}
                for k, v in raw_cache.items():
                    normalized_k = str(Path(k).resolve())
                    if Path(normalized_k).parent == Path(normalized_k) and not normalized_k.endswith(os.sep):
                        normalized_k += os.sep
                    normalized_cache[normalized_k] = v
                self.cache = normalized_cache

                # 3) Merge any cache-only dirs into self.directories, ensuring normalized keys
                for directory_key_raw, data in raw_cache.items(): # Iterate original keys from loaded raw_cache
                    directory = str(Path(directory_key_raw).resolve()) # Use normalized key for processing
                    if Path(directory).parent == Path(directory) and not directory.endswith(os.sep):
                        directory += os.sep
                    
                    if directory not in self.directories:
                        ts = data.get("timestamp", datetime.now().isoformat())
                        self.directories[directory] = ts
                        self.directory_active_status[directory] = True
                
                # 4) Persist the expanded dashboard (with normalized keys)
                self.save_dashboard()

            except Exception as e:
                print(f"[ERROR] Loading cache: {e}")
                return

        # 5) Now build/update a frame for *every* directory we know about (keys are already normalized)
        for directory in sorted(self.directories):
            # update timestamp from cache if available
            if directory in self.cache:
                ts = self.cache[directory].get("timestamp")
                if ts:
                    self.directories[directory] = ts
            # (re)create the UI frame
            self.build_directory_tree(directory)

        print("[LOG] Cache loaded for all known directories.")

    def select_current_document(self):
        file = filedialog.askopenfilename(
            initialdir=self.default_directory or os.getcwd(),
            filetypes=[("Supported", "*.txt *.pdf *.docx *.xls *.html")]
        )
        if not file:
            return
        # Normalize the selected file path
        normalized_file = str(Path(file).resolve())
        
        # 1) Optionally, open it in the OS viewer immediately:
        open_file(normalized_file)

        # 2) Remember this as the current document
        self.current_document = normalized_file
        self.doc_info_title_label.config(text=os.path.basename(normalized_file))

        # 3) Run your full DIG workflow (info + recommendations + dashboard)
        self.dig_document(normalized_file)

        # 4) Play a little chime to signal “select” is done
        winsound.PlaySound('SystemAsterisk', winsound.SND_ALIAS)

    def display_file(self, file):
        # file path is assumed to be normalized here
        base = os.path.splitext(os.path.basename(file))[0]
        first, rest = (base.split(" ", 1) + [""])[:2]
        words = [w for w in base.split() if w.lower() not in COMMON_WORDS]
        parts = words + ([rest] if rest else [])
        self.rest_part = rest

        # --- clear any previous highlight memory ---
        if hasattr(self, "last_selected_parts"):
            for w in self.last_selected_parts:
                if w in self.amplifiers["words"]:
                    self.amplifiers["words"].remove(w)

        # --- 1) Compute the file‐age string (your existing logic) ---
        import datetime as _dt
        mtime = _dt.datetime.fromtimestamp(os.path.getmtime(file))
        now   = _dt.datetime.now()
        age_days = (now - mtime).days
        if age_days == 0:
            age_str = "Today"
        elif age_days == 1:
            age_str = "Yesterday"
        else:
            age_str = f"{age_days} days ago"

        # --- 2) Retrieve highlights based on extension ---
        ext = os.path.splitext(file)[1].lower()
        if ext == ".html":
            highlights = retrieve_html_highlights(file)
        elif ext == ".docx":
            highlights = retrieve_docx_highlights(file)
        elif ext == ".pdf":
            highlights = retrieve_pdf_highlights(file)
        else:
            highlights = []

        # --- 3) Build the intro and full_text exactly as before ---
        if highlights:
            hl_intro = f"\n\nFound {len(highlights)} highlighted passage(s):\n\n"
        else:
            hl_intro = "\n\nNo highlights found.\n\n"

        # assume your original code split the file into `parts` and `rest`

        self.rest_part = rest

        full_text = (
            f"File: {os.path.basename(file)}\n"
            f"Last modified: {mtime.strftime('%Y-%m-%d %H:%M')}"
            f"  ({age_str})\n"
            + hl_intro
            + rest  # or however you combine the non‐highlight text
        )

        # --- 4) Populate the text widget ---
        self.output_text.config(state="normal")
        self.output_text.delete("1.0", tk.END)
        self.output_text.insert(tk.END, full_text)

        # insert the actual highlights in color
        for color_hex, hl_text, style_dict in highlights:
            tag_name = f"hl_{color_hex}_{abs(hash(hl_text))}"
            if tag_name not in self.output_text.tag_names():
                self.output_text.tag_configure(tag_name, **style_dict)
            self.output_text.insert(tk.END, hl_text + "\n\n", tag_name)

        self.output_text.config(state="disabled")

        # --- 5) finish up your dashboard/amplifiers logic ---
        self.amplifiers["words"].extend(parts)
        self.last_selected_parts = parts.copy()
        self.refresh_amp_list()
        self.save_dashboard()
        self.current_document = file

    def show_document_info(self, file, text):
        # Store current context for real-time updates
        self.current_file = file
        self.current_text = text

        # Configuration
        MAX_MATCH_WORDS = 25  # maximum words to display per match
        MAX_MATCHES = 20      # maximum number of matches to display

        import os, re, tkinter as tk
        import tkinter.font as tkfont
        from datetime import datetime

        # Normalize file path and parse publication date
        base = os.path.splitext(os.path.basename(file))[0]
        first_token = base.split()[0] if base else ""
        age_str = ""
        if first_token.isdigit() and len(first_token) == 8:
            try:
                pub_date = datetime.strptime(first_token, "%Y%m%d")
                y, m, d = compute_age(pub_date)
                age_str = f"Published {y} years, {m} months, {d} days ago\n\n"
            except Exception:
                age_str = ""

        # Extract query terms
        query = self.search_var.get().strip()
        phrases, terms = [], []
        for match in re.finditer(r'"([^\"]+)"|(\S+)', query):
            if match.group(1):
                phrases.append(match.group(1))
            else:
                term = match.group(2)
                if term.lower() != 'ayo':
                    terms.append(term)
        tokens = phrases + terms

        # Split text into sentences and collect unique matches
        sentences = re.split(r'(?<=[\.\!?])\s+', text)
        unique_matches = []
        for phrase in phrases:
            for sent in sentences:
                if re.search(re.escape(phrase), sent, re.IGNORECASE) and sent.strip() not in unique_matches:
                    unique_matches.append(sent.strip())
        for term in terms:
            for sent in sentences:
                if re.search(re.escape(term), sent, re.IGNORECASE) and sent.strip() not in unique_matches:
                    unique_matches.append(sent.strip())

        # Limit number of search matches
        unique_matches = unique_matches[:MAX_MATCHES]

        # Build snippet
        snippet = ""
        if unique_matches:
            snippet_lines = []
            for idx, sent in enumerate(unique_matches, 1):
                words = sent.split()
                positions = [i for token in tokens for i, w in enumerate(words)
                             if re.search(re.escape(token), w, re.IGNORECASE)]
                center = positions[0] if positions else 0
                half = MAX_MATCH_WORDS // 2
                start = max(0, center - half)
                end = start + MAX_MATCH_WORDS
                if end > len(words):
                    end = len(words)
                    start = max(0, end - MAX_MATCH_WORDS)
                window = words[start:end]
                display = " ".join(window)
                if start > 0:
                    display = "..." + display
                if end < len(words):
                    display += "..."
                snippet_lines.append(f"{idx}. {display}")
            snippet = "Search matches:\n" + "\n".join(snippet_lines) + "\n\n"

        # Retrieve existing highlights
        ext = os.path.splitext(file)[1].lower()
        if ext == ".html":
            highlights = retrieve_html_highlights(file)
        elif ext == ".docx":
            highlights = retrieve_docx_highlights(file)
        elif ext == ".pdf":
            highlights = retrieve_pdf_highlights(file)
        else:
            highlights = []

        # Highlights intro
        hl_intro = f"Found {len(highlights)} highlighted passage(s):\n\n" if highlights else "No highlights found.\n\n"

        # Prepare text widget
        self.output_text.config(state="normal")
        self.output_text.delete("1.0", tk.END)
        self.output_text.insert(tk.END, snippet + age_str + hl_intro)

        # Base font for document info
        base_font = tkfont.Font(font=self.output_text.cget("font"))

        # Function to build highlight font
        def highlight_font(style_dict):
            f = base_font.copy()
            f.configure(size=self.doc_info_font_size)
            if style_dict.get("bold"):
                f.configure(weight="bold")
            if style_dict.get("italic"):
                f.configure(slant="italic")
            return f

        # Insert highlights
        for color_hex, hl_text, style_dict in highlights:
            tag_name = f"hl_{color_hex}_{abs(hash(hl_text))}"
            if tag_name not in self.output_text.tag_names():
                tag_font = highlight_font(style_dict)
                self.output_text.tag_configure(tag_name, font=tag_font, background=color_hex)
            self.output_text.insert(tk.END, hl_text + "\n\n", tag_name)

        # Original first 500 words
        self.output_text.insert(tk.END, "Original (first 500 words):\n\n")
        words = text.split()
        first_words = " ".join(words[:500])
        self.output_text.insert(tk.END, first_words)

        # Reapply highlights in original snippet
        for color_hex, hl_text, style_dict in highlights:
            tag = f"orig_hl_{color_hex}_{abs(hash(hl_text))}"
            if tag not in self.output_text.tag_names():
                tag_font = highlight_font(style_dict)
                self.output_text.tag_configure(tag, font=tag_font, background=color_hex)
            start = "1.0"
            while True:
                idx = self.output_text.search(hl_text, start, tk.END)
                if not idx:
                    break
                end_idx = f"{idx}+{len(hl_text)}c"
                self.output_text.tag_add(tag, idx, end_idx)
                start = end_idx

        # Bold tag for search terms
        if "search_bold" not in self.output_text.tag_names():
            bold_font = base_font.copy()
            bold_font.configure(size=self.doc_info_font_size, weight="bold")
            self.output_text.tag_configure("search_bold", font=bold_font)

        # Apply bold to tokens
        for token in tokens:
            start = "1.0"
            while True:
                idx = self.output_text.search(token, start, tk.END, nocase=1)
                if not idx:
                    break
                end_idx = f"{idx}+{len(token)}c"
                self.output_text.tag_add("search_bold", idx, end_idx)
                start = end_idx

        self.output_text.config(state="disabled")

    def on_font_size_change(self):
        # Update stored font size
        new_size = self.font_size_var.get()
        self.doc_info_font_size = new_size
        # Update snippet font
        current_font = tkfont.Font(font=self.output_text.cget("font"))
        current_font.configure(size=new_size)
        self.output_text.configure(font=current_font)
        # Refresh entire display to reapply highlights and search bold with new size
        if hasattr(self, 'current_file') and hasattr(self, 'current_text'):
            self.show_document_info(self.current_file, self.current_text)
        # Persist setting
        self.save_dashboard()

    def extract_text(self, path):
        # path is assumed to be normalized here
        ext = os.path.splitext(path)[1].lower()
        if ext == ".txt":
            return Path(path).read_text(errors="ignore")
        elif ext == ".pdf":
            with fitz.open(path) as doc:
                return " ".join(page.get_text() for page in doc)
        elif ext == ".docx":
            return " ".join(p.text for p in docx.Document(path).paragraphs)
        elif ext == ".xls":
            book = xlrd.open_workbook(path)
            return " ".join(str(cell.value) for sheet in book.sheets() for row in range(sheet.nrows) for cell in sheet.row(row))
        elif ext == ".html":
            return BeautifulSoup(Path(path).read_text(errors="ignore"), "html.parser").get_text()
        return ""

    def recommend_similar_files(self, current_path):
        print(f"[LOG] Analyzing: {current_path}")
        current_text = self.extract_text(current_path)

        # Gather all candidate texts
        candidate_texts = []
        file_map = []
        for directory, data in self.cache.items():
            if not self.directory_active_status.get(directory, True):
                continue
            candidate_texts.extend(data["texts"])
            file_map.extend(data["files"])
        all_texts = [current_text] + candidate_texts

        # 1) Vectorize everything
        vect = TfidfVectorizer()
        tfidf = vect.fit_transform(all_texts)            # shape: (1 + num_candidates) × V
        vocab = vect.vocabulary_                         # word → column index

        # 2) **Amplify** the query’s TF‑IDF for each amp‑word
        F = self.amplifiers["factor"]
        for w in self.amplifiers["words"]:
            key = w.lower()
            if key in vocab:
                col = vocab[key]
                # multiply only the query vector (row 0)
                tfidf[0, col] *= F

        # 3) Compute boosted similarities
        similarities = cosine_similarity(tfidf[0:1], tfidf[1:]).ravel()

        # 4) Apply silencers and length penalties as before
        WORD_COUNT_THRESHOLD = 3000
        LENGTH_FACTOR        = 0.8
        adjusted_scores = []
        rest = getattr(self, "rest_part", "").lower()

        for sim, candidate in zip(similarities, candidate_texts):
            adj = sim

            # rest‑piece super‑boost
            if rest and rest in candidate.lower():
                adj *= (self.amplifiers["factor"] * 300)
            else:
                # any other amp‑word (normal boost)
                for word in self.amplifiers["words"]:
                    if word.lower() in candidate.lower():
                        adj *= self.amplifiers["factor"]
                        break

            # silencers
            for word in self.silencers["words"]:
                if word.lower() in candidate.lower():
                    adj *= self.silencers["factor"]
                    break

            # length penalty
            word_count = len(candidate.split())
            if word_count > WORD_COUNT_THRESHOLD:
                adj *= LENGTH_FACTOR

            adjusted_scores.append(adj)

        # 5) Take top‑12 and emit shortcuts
        ranked = sorted(zip(adjusted_scores, file_map),
                        key=lambda x: x[0], reverse=True)[:12]
        self.clear_shortcut_buttons()
        self.clear_shortcuts()
        for score, path in ranked:
            name = f"{score:.4f}_{os.path.basename(path)}"
            shortcut_path = SHORTCUT_DIR / (name + ".lnk")
            try:
                winshell.CreateShortcut(
                    Path=str(shortcut_path),
                    Target=path,
                    Icon=(path, 0),
                    Description="Similar Document"
                )
            except Exception as e:
                print(f"[ERROR] Failed to create shortcut for {path}: {e}")
            print(f"[LOG] Recommendation {score:.4f}: {path}")

        self.recommended_docs = ranked
        self.refresh_recommendation_buttons()

    def clear_shortcuts(self):
        for file in SHORTCUT_DIR.iterdir():
            if file.suffix == ".lnk" and "_" in file.stem:
                try:
                    file.unlink()
                except Exception as e:
                    print(f"[ERROR] Error deleting shortcut {file}: {e}")

    def refresh_toggle_buttons(self):
        # 1) Wipe out any old buttons/frames
        for frm in self.toggle_frames.values():
            frm.destroy()
        self.toggle_buttons.clear()
        self.toggle_frames.clear()

        # 2) Rebuild in sorted(self.directories) order (keys are normalized)
        for directory in sorted(self.directories):
            is_active = self.directory_active_status.get(directory, True)
            
            # Use appropriate display name for the toggle button
            display_name = os.path.basename(directory)
            if not display_name and Path(directory).is_dir() and Path(directory).drive: # It's a root drive like C:\
                display_name = directory # Use the full path for display if it's a root drive
            elif not display_name: # Fallback for non-drive roots (e.g., '/') or other edge cases
                display_name = directory

            frm = ttk.Frame(self.toggle_btns_frame)
            frm.pack(fill="x", pady=2)
            btn = tk.Button(frm,
                text=display_name, # Use display_name here
                bg="green" if is_active else "red",
                fg="white",
                relief="raised",
                command=lambda d=directory: self._on_sidebar_toggle(d) # Pass normalized directory
            )
            btn.pack(side="left", fill="x", expand=True)

            erase = tk.Button(frm,
                text="Erase",
                bg="#ff4444", fg="white", relief="raised",
                command=lambda d=directory: self._on_erase_sidebar_directory(d) # Pass normalized directory
            )
            erase.pack(side="right", padx=(5,0))

            self.toggle_frames[directory]  = frm
            self.toggle_buttons[directory] = btn

    def open_tree_to_path(self, target_path):
        import os

        # 1) Collapse every top‑level node so we start from a clean slate
        for iid in self.dir_tree.get_children(""):
            self.dir_tree.item(iid, open=False)

        # 2) pick deepest cached root
        candidates = [d for d in self.directories if target_path.startswith(d)]
        if not candidates:
            return
        root_dir = max(candidates, key=len)

        # 3) find its IID
        root_iid = None
        for iid in self.dir_tree.get_children(""):
            if self.dir_tree.item(iid, "values")[0] == root_dir:
                root_iid = iid
                break
        if root_iid is None:
            return

        # 4) break into folder segments + filename
        rel = os.path.relpath(target_path, root_dir)
        parts = rel.split(os.sep)
        folder_segments, filename = parts[:-1], parts[-1]

        # 5) walk down – expand each folder programmatically
        current = root_iid
        for seg in folder_segments:
            # open & populate this node
            self._expand_and_populate(current)

            # find the matching child
            for child in self.dir_tree.get_children(current):
                if self.dir_tree.item(child, "text") == seg:
                    current = child
                    break
            else:
                # segment not found
                return

        # 6) finally expand the leaf folder
        self._expand_and_populate(current)

        # 7) select the file and scroll into view
        for leaf in self.dir_tree.get_children(current):
            if self.dir_tree.item(leaf, "text") == filename:
                self.dir_tree.selection_set(leaf)
                self.dir_tree.see(leaf)
                break

    def _insert_real_children(self, iid):
        """
        Given a folder‐node IID, list its filesystem directory
        and insert each entry as a child in the Treeview,
        adding a dummy "…" for subfolders.
        """
        import os

        # determine the filesystem path for this node
        dirpath = self.dir_tree.item(iid, "values")[0]

        try:
            entries = sorted(os.listdir(dirpath))
        except Exception:
            return

        for name in entries:
            fullpath = os.path.join(dirpath, name)
            # insert the child node
            child_iid = self.dir_tree.insert(iid, "end",
                                             text=name,
                                             values=(fullpath,))
            # if it's a directory, add the dummy so it can be expanded later
            if os.path.isdir(fullpath):
                self.dir_tree.insert(child_iid, "end", text="…")
            if self.view_mode.get() == "Saved" and fullpath not in self.saved_files:
                continue

    def dir_contains_saved(self, directory):
        return any(fp.startswith(directory) for fp in self.saved_files)

    def _expand_and_populate(self, iid):
        # open the node
        self.dir_tree.item(iid, open=True)

        # if it only has the dummy placeholder, replace it
        children = self.dir_tree.get_children(iid)
        if len(children) == 1 and self.dir_tree.item(children[0], "text") == "…":
            # delete the dummy
            self.dir_tree.delete(children[0])
            # insert the real children
            self._insert_real_children(iid)
            # update UI
            self.dir_tree.update_idletasks()

    def _on_tree_expand_programmatic(self, iid):
        children = self.dir_tree.get_children(iid)
        if len(children) == 1 and self.dir_tree.item(children[0], "text") == "…":
            self.dir_tree.delete(children[0])
            # Normalize fullpath when retrieving from item values
            fullpath = str(Path(self.dir_tree.item(iid, "values")[0]).resolve())

            try:
                for name in sorted(os.listdir(fullpath)):
                    path = os.path.join(fullpath, name)
                    # Normalize paths before inserting into treeview values
                    normalized_path = str(Path(path).resolve())
                    if os.path.isdir(path):
                        cid = self.dir_tree.insert(iid, "end", text=name, values=(normalized_path,))
                        self.dir_tree.insert(cid, "end", text="…")
                        tag = "active" if self.directory_active_status.get(normalized_path, True) else "inactive" # Use normalized path for status lookup
                        self.dir_tree.item(cid, tags=(tag,))
                    else:
                        fid = self.dir_tree.insert(iid, "end", text=name, values=(normalized_path,))
            except (PermissionError, FileNotFoundError):
                pass

    def _display_recommendation_buttons(self, ranked):
        self.clear_shortcut_buttons()
        if ranked:
            for score, path in ranked: # path is already normalized
                frame = tk.Frame(self.shortcut_buttons_frame, pady=2)
                frame.pack(fill="x", anchor="w", padx=5)
				
                btn_frame = tk.Frame(frame)
                btn_frame.pack(side="top", anchor="w", padx=2)
                btn_open = tk.Button(btn_frame, text="OPEN", command=lambda p=path: open_file(p), width=6, height=1, font=("Helvetica", 8), bg="#c8e6c9")
                btn_open.pack(side="left", padx=(0,2))
                btn_dig = tk.Button(btn_frame, text="DIG", command=lambda p=path: self.dig_document(p), width=6, height=1, font=("Helvetica", 8), bg="#bbdefb")
                btn_dig.pack(side="left", padx=(2,2))
                # Pass normalized path to open_file
                btn_folder = tk.Button(btn_frame, text="FOLDER", command=lambda p=path: open_file(os.path.dirname(p)), width=6, height=1, font=("Helvetica", 8), bg="#ffe0b2")
                btn_folder.pack(side="left", padx=(2,2))
                btn_move = tk.Button(btn_frame, text="MOVE", command=lambda p=path: self.move_selected_file(p), width=6, height=1, font=("Helvetica", 8), bg="#ffcdd2")
                btn_move.pack(side="right", padx=(2,0))
				# new: score right after MOVE
                score_lbl = tk.Label(btn_frame, text=f"{score:.4f}", font=("Helvetica", 10), anchor="w")
                score_lbl.pack(side="left", padx=(2,0))
			
                label = tk.Label(frame, text=os.path.basename(path), font=self._rec_font_normal, anchor="w")
                if self.current_active_file == path:
                    # re‑apply bold font when rebuilding
                    label.config(font=self._rec_font_bold)
                    self.current_active_label = label
                label.pack(fill="x", side="top", padx=2, pady=(0,2))

                # single-click still updates the about-pane
                label.bind("<Button-1>", lambda event, p=path: self.update_about_document(event, p))
                # double-click opens the file in the tree
                label.bind("<Double-1>", lambda event, p=path: self.open_tree_to_path(p))

    def _on_erase_sidebar_directory(self, dir_path):
        # 1) erase from your data
        self.erase_directory(dir_path)
        # 2) This part is handled by refresh_toggle_buttons after erase_directory
        # self.toggle_frames[dir_path].destroy() and del self.toggle_frames[dir_path] are not needed here.
        # 3) Also delete the root node from the tree (handled by erase_directory)
        
    def move_selected_file(self, target_path):
        """Moves the currently selected document to the directory of the target recommended document."""
        # Validate source
        src = self.current_document # This path is already normalized
        if not src or not os.path.isfile(src):
            messagebox.showerror("Error", "No valid file selected to move.")
            return
        # Determine destination - target_path is already normalized
        dest_dir = os.path.dirname(target_path)
        if not os.path.isdir(dest_dir):
            try:
                os.makedirs(dest_dir, exist_ok=True)
            except Exception as e:
                messagebox.showerror("Error", f"Cannot create destination directory: {e}")
                return
        dest_path = os.path.join(dest_dir, os.path.basename(src))
        # Attempt move
        try:
            shutil.move(src, dest_path)
        except shutil.Error as e:
            messagebox.showerror("Move Error", f"Failed to move file (shutil error): {e}")
            return
        except Exception as e:
            messagebox.showerror("Move Error", f"Unexpected error moving file: {e}")
            return
        self.output_text.config(state="normal")
        self.output_text.delete("1.0", tk.END)
        self.output_text.insert(tk.END, f"Document moved to {dest_path}")
        self.output_text.config(state="disabled")

	
    def refresh_recommendation_buttons(self):
        if self.recommended_docs:
            if self.sort_by_filename.get():
                # Paths in recommended_docs are already normalized
                ranked = sorted(self.recommended_docs, key=lambda x: os.path.basename(x[1]).lower(), reverse=True)
            else:
                ranked = sorted(self.recommended_docs, key=lambda x: x[0], reverse=True)
            self._display_recommendation_buttons(ranked)

    def clear_shortcut_buttons(self):
        for widget in self.shortcut_buttons_frame.winfo_children():
            widget.destroy()

    def dig_document(self, path):
        self.ensure_recommendations_visible()
	
        # path is assumed to be normalized here
        # — REMOVE last split parts —
        if hasattr(self, "last_selected_parts"):
            for w in self.last_selected_parts:
                if w in self.amplifiers["words"]:
                    self.amplifiers["words"].remove(w)

        # — SPLIT filename into all words, filter out COMMON_WORDS —
        base = os.path.splitext(os.path.basename(path))[0]
        first, rest = (base.split(" ", 1) + [""])[:2]
        words = [w for w in base.split() if w.lower() not in COMMON_WORDS]
        parts = words + ([rest] if rest else [])
        self.rest_part = rest

        # — ADD those parts to amplifiers in one go —
        self.amplifiers["words"].extend(parts)

        # — REMEMBER them for next time —
        self.last_selected_parts = parts.copy()

        # — REFRESH UI & SAVE —
        self.refresh_amp_list()
        self.save_dashboard()
	
        self.doc_info_title_label.config(text=os.path.basename(path))
        print(f"[LOG] DIG initiated for: {path}")
        self.show_document_info(path, self.extract_text(path))
        self.recommend_similar_files(path)

        # Finally, play a chime to signal “dig” has completed
        winsound.PlaySound('SystemAsterisk', winsound.SND_ALIAS)

    def update_about_document(self, event, file_path):
        # file_path is assumed to be normalized here
        if self.current_active_label is not None and self.current_active_label.winfo_exists():
            self.current_active_label.config(font=self._rec_font_normal)
        else:
            self.current_active_label = None
        event.widget.config(font=self._rec_font_bold)
        self.current_active_label = event.widget
        self.current_active_file = file_path
        text = self.extract_text(file_path)
        self.show_document_info(file_path, text)
        print(f"[LOG] Updated About panel for: {file_path}")

if __name__ == "__main__":
    root = tk.Tk()
    root.rowconfigure(0, weight=1)
    root.columnconfigure(0, weight=1)
    app = DocumentRecommenderApp(root)
    root.mainloop()