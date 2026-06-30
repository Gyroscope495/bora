# --- bora.pyw ---
import os
import sys
import tkinter as tk

# --- 1. INSTANT SPLASH SCREEN ---
# We do this before heavy imports so the logo appears the millisecond the user clicks the app.
if __name__ == "__main__":
    try:
        import ctypes
        ctypes.windll.user32.ShowWindow(ctypes.windll.kernel32.GetConsoleWindow(), 0)
    except Exception: pass
    
    root = tk.Tk()
    root.withdraw()
    
    from startup import show_startup_logo
    splash, progress_var, splash_status = show_startup_logo(root)
    root.update()
    
    def update_splash(msg, val=None):
        splash_status.config(text=msg)
        if val is not None:
            progress_var.set(val)
        root.update()
        
    update_splash("Loading core libraries...", 5)

import json
import shutil

from tkinter import filedialog, messagebox, ttk
from tkinter.ttk import Progressbar
import time
from datetime import datetime
from pathlib import Path
import subprocess
import tkinter.font as tkfont
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime as _dt
import re
import calendar
import random
from tkinter import colorchooser

# Local module imports
from src import recommender
from src import documentinfo
from src import search
from src import cached_saved_recent_button
from src.zoom_view import apply_zoom_view
from src import cache_generator
from src.cache_generator import SUPPORTED_EXTENSIONS
from src import year_lookback

# --- Optional chime, safe on non-Windows ---
try:
    import winsound
    def _play_chime():
        winsound.PlaySound('SystemAsterisk', winsound.SND_ALIAS)
except Exception:
    def _play_chime():
        pass

# --- Hide console window on Windows (if present) ---
def hide_console_window():
    """Hides the console window on Windows systems. No-op on other OS."""
    try:
        import ctypes
        kernel32 = ctypes.windll.kernel32
        user32 = ctypes.windll.user32
        get_console = kernel32.GetConsoleWindow
        show_window = user32.ShowWindow
        SW_HIDE = 0
        hwnd = get_console()
        if hwnd:
            show_window(hwnd, SW_HIDE)
    except Exception:
        pass

def maximize_window(root):
    """Best-effort maximize across platforms without forcing fullscreen."""
    try:
        root.state('zoomed')  # Windows (and some X11)
        return
    except Exception:
        pass
    try:
        root.attributes('-zoomed', True)
        return
    except Exception:
        pass
    try:
        sw, sh = root.winfo_screenwidth(), root.winfo_screenheight()
        root.geometry(f"{sw}x{sh}+0+0")
    except Exception:
        pass

def set_app_icon(root):
    """Sets the application icon, prioritizing .ico for Windows."""
    script_dir = Path(__file__).parent
    
    # Prioritize .ico for Windows, fallback to .png/.gif
    ico_path = script_dir / "assets" / "logo.ico"
    png_path = script_dir / "assets" / "logo.png"
    gif_path = script_dir / "assets" / "logo.gif"

    try:
        if os.name == "nt" and ico_path.exists():
            root.iconbitmap(str(ico_path))
        elif png_path.exists():
            img = tk.PhotoImage(file=str(png_path))
            root.icon_image_ref = img  # Keep reference
            root.iconphoto(True, img)
        elif gif_path.exists():
            img = tk.PhotoImage(file=str(gif_path))
            root.icon_image_ref = img  # Keep reference
            root.iconphoto(True, img)
    except Exception as e:
        print(f"[WARN] Could not set application icon: {e}")


# Set up paths
# Set up paths
if getattr(sys, 'frozen', False):
    # We are running as a compiled .exe
    BASE_DIR = Path(sys.executable).parent
else:
    # We are running as a .py script
    BASE_DIR = Path(__file__).parent

CACHE_DIR = BASE_DIR / "cache"
SHORTCUT_DIR = BASE_DIR
CACHE_FILE = CACHE_DIR / "cache.json"
DASHBOARD_FILE = CACHE_DIR / "dashboard.json"

CACHE_DIR.mkdir(exist_ok=True)
CACHE_FILE = CACHE_DIR / "cache.json"
DASHBOARD_FILE = CACHE_DIR / "dashboard.json"

CACHE_DIR.mkdir(exist_ok=True)

# Define default width for the cached directories sidebar
DEFAULT_SIDEBAR_WIDTH = 150

# define words to ignore (stop-words)
ignore_path = BASE_DIR / "ignore_words.txt"
COMMON_WORDS = set()

if ignore_path.exists():
    with open(ignore_path, "r", encoding="utf-8") as f:
        COMMON_WORDS = {line.strip().lower() for line in f if line.strip()}
else:
    # Fallback and auto-create the file for future use
    COMMON_WORDS = {
        "a", "an", "the", "and", "or", "of", "in", "on", "for", "to", "with", "at", 
        "by", "from", "up", "out", "over", "is", "isn't", "us", "à", "o", "e", "de", 
        "do", "da", "em", "um", "uma", "que", "com", "por", "para", "has", "billion", 
        "part", "parts", "1", "2", "3", "4", "5", "6", "7", "8", "9", "0", "-", "$", 
        "&", "–", "'", "_", "first", "two", "three", "four", "five", "six", "seven", 
        "eight", "nine", "ten", "um", "dois", "três", "quatro", "cinco", "seis", 
        "sete", "oito", "nove", "dez", "I", "you", "we", "them", "eu", "ele", "ela", 
        "eles", "elas", "isso", "isto", "aquilo", "how", "what", "why", "where", 
        "would", "here", "there", "when", "who", "which", "do", "does", "did", "done",
        "solve", "as", "big", "small", "lull", "agents", "agent", "after", "before",
        "agora", "os", "é", "mais", "really", "very", "much", "many", "most", "some", 
        "few", "lot", "lots", "realmente", "muito", "muitos", "a maioria", "alguns", 
        "poucos", "bastante", "edge", "guardian", "casa", "agente", "agentes", "ein", ","
    }
    try:
        with open(ignore_path, "w", encoding="utf-8") as f:
            f.write("\n".join(sorted(list(COMMON_WORDS))))
    except Exception as e:
        print(f"[WARN] Could not auto-create ignore_words.txt: {e}")

AMPLIFIERS_FILE = BASE_DIR / "amplifiers.txt"

def load_amplifiers_from_txt():
    """Reads amplifiers from the text file. Creates it with defaults if missing."""
    if AMPLIFIERS_FILE.exists():
        with open(AMPLIFIERS_FILE, "r", encoding="utf-8") as f:
            return [line.strip() for line in f if line.strip()]
    else:
        # Fallback defaults if the file was deleted or it's the first run
        default_words = ["pensascaramouche", "fuchspensamentos", "mentescrita", "§", "áudio", "criatividade"]
        try:
            with open(AMPLIFIERS_FILE, "w", encoding="utf-8") as f:
                f.write("\n".join(default_words))
        except Exception as e:
            print(f"[WARN] Could not auto-create amplifiers.txt: {e}")
        return default_words

def save_amplifiers_to_txt(words_list):
    """Saves the dashboard's active amplifiers back to the text file."""
    try:
        with open(AMPLIFIERS_FILE, "w", encoding="utf-8") as f:
            f.write("\n".join(words_list))
    except Exception as e:
        print(f"[ERROR] Could not save amplifiers.txt: {e}")

def open_file(path):
    """Opens a file or folder using the OS default application without spawning a console window."""
    try:
        if os.name == "nt":
            os.startfile(path)
        elif sys.platform == "darwin":
            subprocess.run(["open", path], check=False)
        else:
            subprocess.run(["xdg-open", path], check=False)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to open: {e}")

def check_for_color_highlights(path):
    import fitz  # Deferred import
    ext = os.path.splitext(path)[1].lower()
    if ext != ".pdf":
        return "N/A"
    try:
        doc = fitz.open(path)
        for page in doc:
            annots = page.annots()
            if annots is None:
                continue
            for annot in annots:
                if annot.type[0] == 8:
                    return "yes"
        return "no"
    except Exception as e:
        return f"error: {e}"

def retrieve_color_highlights(path):
    import fitz  # Deferred import
    ext = os.path.splitext(path)[1].lower()
    if ext != ".pdf":
        return []
    highlights = []
    try:
        doc = fitz.open(path)
        for page in doc:
            annots = page.annots()
            if annots is None:
                continue
            for annot in annots:
                if annot.type[0] == 8:  # Highlight annotation
                    highlight_text = annot.info.get("content", "").strip()
                    if not highlight_text:
                        r = annot.rect
                        highlight_text = page.get_text("text", clip=r).strip()
                    colors = annot.colors
                    if "fill" in colors and colors["fill"]:
                        color_tuple = colors["fill"]
                    elif "stroke" in colors and colors["stroke"]:
                        color_tuple = colors["stroke"]
                    else:
                        color_tuple = (0, 0, 0)
                    try:
                        r_val, g_val, b_val = color_tuple
                        if r_val <= 1:
                            r_int = int(r_val * 255)
                            g_int = int(g_val * 255)
                            b_int = int(b_val * 255)
                        else:
                            r_int, g_int, b_int = int(r_val), int(g_val), int(b_val)
                        color_hex = f"#{r_int:02x}{g_int:02x}{b_int:02x}"
                    except Exception:
                        color_hex = "#000000"
                    highlights.append((color_hex, highlight_text))
        return highlights
    except Exception as e:
        print(f"[ERROR] retrieve_color_highlights: {e}")
        return []

def build_font_from_style(style_dict, base_font=("Courier", 10)):
    family = style_dict.get("font-family", base_font[0])
    size_str = style_dict.get("font-size", None)
    try:
        size = int(size_str.replace("px", "").strip()) if size_str else base_font[1]
    except Exception:
        size = base_font[1]
    weight = style_dict.get("font-weight", "").lower()
    font_weight = "bold" if "bold" in weight else "normal"
    slant_val = style_dict.get("font-style", "").lower()
    font_slant = "italic" if "italic" in slant_val else "roman"
    return tkfont.Font(family=family, size=size, weight=font_weight, slant=font_slant)

class DocumentRecommenderApp:
    def __init__(self, root):
        self.root = root
        self.root.title("bloco v7")
        self.root.geometry("800x700")
        self.root.rowconfigure(0, weight=1)
        self.root.columnconfigure(0, weight=1)
        # Fix for Windows themes ignoring Treeview tag colors
        style = ttk.Style()
        style.map('Treeview', foreground=style.map('Treeview', 'foreground'), background=style.map('Treeview', 'background'))
        self.sidebar_width = DEFAULT_SIDEBAR_WIDTH

        self.doc_info_font_size = 10
        self.presets = {
            "1": {"name": "1", "state": {}, "color": "#e0e0e0"},
            "2": {"name": "2", "state": {}, "color": "#e0e0e0"},
            "3": {"name": "3", "state": {}, "color": "#e0e0e0"}
        }
        self.active_preset = None  # Tracks the currently active preset
        self.max_recommendations = 20

        self.directories = {}
        self.directory_colors = {} # <--- NEW
        self.cache = {}
        self.directory_frames = {}
        self.directory_active_status = {}
        self.toggle_buttons = {}
        self.toggle_frames = {}
        self.default_directory_used = False

        self.silencers = {"words": [], "factor": 1.0}
        self.length_penalty = {"threshold": 3000, "factor": 0.9}
        self.recent_timespan_hours = 24
        self.default_directory = ""
        self.saved_files = set()
        self.load_dashboard()

        self.filepath_var = tk.StringVar(value="")
        self.current_document = None
        self.recommended_docs = []
        self.sort_by_filename = tk.BooleanVar(value=False)
        self.amplify_results = tk.BooleanVar(value=False)
        self.current_active_file = None
        self.current_active_label = None
        
        # For single-file cache updates
        self.last_opened_for_editing = None
        self.last_opened_mtime = None
        self.root.protocol("WM_DELETE_WINDOW", self._on_closing)


        self._rec_font_normal = ("Helvetica", 10, "normal")
        self._rec_font_bold   = ("Helvetica", 10, "bold")

        self.setup_scrollable_frame()
        self.create_widgets()
        # Removed self.load_cache() from here so the splash screen can track it
    
    def _highlight_and_flash_search_terms(self):
        query = self.search_var.get().strip()
        if not query:
            return

        self.output_text.config(state="normal")
        
        # 1. Permanent highlight tag for search terms (light yellow)
        self.output_text.tag_configure("search_term", background="#ffff99", foreground="black")
        
        # 2. Temporary flash tag (bright orange)
        self.output_text.tag_configure("search_flash", background="#ff8c00", foreground="white")
        
        words = [w for w in query.lower().split() if w not in COMMON_WORDS]
        
        for word in words:
            start_idx = "1.0"
            while True:
                start_idx = self.output_text.search(word, start_idx, nocase=True, stopindex=tk.END)
                if not start_idx:
                    break
                end_idx = f"{start_idx}+{len(word)}c"
                
                # Apply both tags
                self.output_text.tag_add("search_term", start_idx, end_idx)
                self.output_text.tag_add("search_flash", start_idx, end_idx)
                
                start_idx = end_idx
        
        self.output_text.config(state="disabled")
        
        # Remove the flash tag after 1.5 seconds, leaving only the permanent highlight
        self.root.after(1500, lambda: self.output_text.tag_remove("search_flash", "1.0", tk.END))

    def change_preset_color(self, p_id):
        """Opens a color picker to change the active color of a preset."""
        current_color = self.presets[p_id].get("color", "#e0e0e0")
        new_color = colorchooser.askcolor(initialcolor=current_color, title=f"Pick active color for '{self.presets[p_id]['name']}'")[1]
        
        if new_color:
            self.presets[p_id]["color"] = new_color
            self.save_dashboard()
            # Immediately update the UI if this is the currently active preset
            if getattr(self, "active_preset", None) == p_id:
                self._update_preset_buttons_ui()

    def _update_preset_buttons_ui(self):
        """Applies the custom color to the active preset and resets the others."""
        for pid, btn in self.preset_buttons.items():
            if pid == getattr(self, "active_preset", None):
                bg_color = self.presets[pid].get("color", "#e0e0e0")
                fg_color = self._get_contrast_color(bg_color)
            else:
                bg_color = "#e0e0e0"
                fg_color = "black"
            btn.config(bg=bg_color, fg=fg_color)

    def _clear_active_preset(self):
        """Resets the active preset state if a manual change is made."""
        if getattr(self, "active_preset", None) is not None:
            self.active_preset = None
            self._update_preset_buttons_ui()

    def _i_have_thoughts(self):
            """Creates a new .docx file based on the selected file and opens it."""
            import docx  # Ensure docx is available
        
            selected_item = self.dir_tree.focus()
            if not selected_item:
                return
            
            original_path = self._node_path(selected_item)
            if not os.path.isfile(original_path):
                messagebox.showwarning("Selection", "Please select a file to attach thoughts to.")
                return

            # 1. Prepare paths and names
            directory = os.path.dirname(original_path)
            base_name = os.path.splitext(os.path.basename(original_path))[0]
            today_str = datetime.now().strftime("%Y%m%d")
        
            new_filename = f"{base_name} {today_str} pensascaramouche.docx"
            new_full_path = os.path.join(directory, new_filename)

            # 2. Check if file already exists to avoid overwriting
            if os.path.exists(new_full_path):
                # If it exists, just open it
                open_file(new_full_path)
                return

            # 3. Create the document
            try:
                doc = docx.Document()
                # Optional: Add a title or reference to the original file inside the doc
                doc.add_paragraph(f"pensascaramouche sobre: {base_name}")
                doc.add_paragraph(f"Created on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                doc.add_paragraph("-" * 20)
            
                doc.save(new_full_path)
            
                # 4. Open the document
                open_file(new_full_path)
            
                if hasattr(self, "status_label"):
                    self.status_label.config(text=f"Created thoughts: {new_filename}")
                
            except Exception as e:
                messagebox.showerror("Error", f"Could not create Word document:\n{e}")
    
    def _generate_random_color(self):
        """Generates a random pastel hex color so black text remains readable."""
        r = random.randint(200, 255)
        g = random.randint(200, 255)
        b = random.randint(200, 255)
        return f"#{r:02x}{g:02x}{b:02x}"

    def _change_dir_color(self, directory):
        """Opens a color chooser to let the user change a directory's color."""
        current_color = self.directory_colors.get(directory, "#ffffff")
        new_color = colorchooser.askcolor(initialcolor=current_color, title=f"Pick color for {os.path.basename(directory)}")[1]
        if new_color:
            self.directory_colors[directory] = new_color
            self.save_dashboard()
            self.refresh_toggle_buttons()
            self.build_directory_tree(directory)  # Re-renders tree elements to apply the new color
            self.refresh_recommendation_buttons()
    
    def _get_depth_color(self, base_color_hex, depth):
        """Lightens a hex color based on depth. Each level adds 15% white to create a cascading effect."""
        if not base_color_hex or not base_color_hex.startswith('#') or len(base_color_hex) != 7:
            return base_color_hex
        try:
            r = int(base_color_hex[1:3], 16)
            g = int(base_color_hex[3:5], 16)
            b = int(base_color_hex[5:7], 16)
            
            # Factor caps at 85% white so it never becomes completely invisible
            factor = min(depth * 0.15, 0.85)
            
            r = int(r + (255 - r) * factor)
            g = int(g + (255 - g) * factor)
            b = int(b + (255 - b) * factor)
            
            return f"#{r:02x}{g:02x}{b:02x}"
        except Exception:
            return base_color_hex

    def _get_contrast_color(self, hex_color):
        """Returns 'white' or 'black' depending on the background's luminance."""
        if not hex_color or not hex_color.startswith('#') or len(hex_color) != 7:
            return "black"
        try:
            r = int(hex_color[1:3], 16)
            g = int(hex_color[3:5], 16)
            b = int(hex_color[5:7], 16)
            # HSP equation for perceived brightness
            luminance = 0.299 * r + 0.587 * g + 0.114 * b
            return "white" if luminance < 140 else "black"
        except Exception:
            return "black"

    def find_in_tree(self, target_path):
        """Forces 'Cached tree' view, clears search, and opens the tree to the path."""
        needs_refresh = False
        
        if self.view_mode.get() != "Cached tree":
            self.view_mode.set("Cached tree")
            self.view_mode_button.config(text="Cached tree", bg="#ccffcc")
            needs_refresh = True
            
        if self.search_var.get().strip() != "":
            self.search_var.set("")
            needs_refresh = True
            
        if needs_refresh:
            cached_saved_recent_button.on_view_mode_change(self)
            
        # Give the UI a split second to redraw the full tree before trying to open the path
        self.root.after(100, lambda: self.open_tree_to_path(target_path))

    def _get_tree_depth(self, iid):
        """Calculates how deep a node is in the Treeview."""
        depth = 0
        parent = self.dir_tree.parent(iid)
        while parent != "":
            depth += 1
            parent = self.dir_tree.parent(parent)
        return depth

    def _get_tree_root_path(self, iid):
        """Traces back to the root node to find the base directory path."""
        current = iid
        parent = self.dir_tree.parent(current)
        while parent != "":
            current = parent
            parent = self.dir_tree.parent(current)
        values = self.dir_tree.item(current, "values")
        return values[0] if values else ""

    def setup_scrollable_frame(self):
        self.canvas = tk.Canvas(self.root)
        self.v_scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=self.v_scrollbar.set)
        self.v_scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)

        self.scroll_container = tk.Frame(self.canvas)
        self.paned = ttk.Panedwindow(self.scroll_container, orient="horizontal")
        self.paned.pack(fill="both", expand=True)

        self.sidebar = ttk.Frame(self.paned, width=self.sidebar_width, relief="sunken")
        self.sidebar.configure(width=self.sidebar_width)
        self.paned.add(self.sidebar, weight=0)

        self.right_pane = ttk.Frame(self.paned)
        self.paned.add(self.right_pane, weight=1)
        
        # --- ROBUST SASH POSITIONING ---
        def _apply_sash_pos(event=None):
            """Applies the sash position exactly when the paned window appears on screen."""
            self.root.update_idletasks() # Ensure geometry is calculated
            self.paned.sashpos(0, self.sidebar_width)
            # Unbind so it only happens once during startup
            self.paned.unbind("<Map>")

        # Wait until the paned window actually appears (after the splash screen hides)
        self.paned.bind("<Map>", _apply_sash_pos)

        def on_sash_drag(event=None):
            # Trigger the save slightly after the mouse lets go
            self.root.after(50, _save_sash_width)
            
        def _save_sash_width():
            try:
                new_w = self.paned.sashpos(0)
                if new_w > 0:
                    self.sidebar_width = new_w
                    self.save_dashboard()
            except Exception:
                pass

        # ONLY save the width when the user RELEASES the mouse button.
        self.paned.bind("<ButtonRelease-1>", on_sash_drag)
        # -------------------------------

        self.canvas_window = self.canvas.create_window((0, 0), window=self.scroll_container, anchor="nw")
        self.scroll_container.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.bind("<Configure>", self.on_canvas_configure)

    def on_font_size_change(self):
        new_size = self.font_size_var.get()
        self.doc_info_font_size = new_size
        current_font = tkfont.Font(font=self.output_text.cget("font"))
        current_font.configure(size=new_size)
        self.output_text.configure(font=current_font)
        self.save_dashboard()

    def on_canvas_configure(self, event):
        self.canvas.itemconfig(self.canvas_window, width=event.width, height=event.height)

    # ---------- dialogs ----------
    def _open_font_dialog(self):
        win = tk.Toplevel(self.root)
        win.title("Font")
        win.transient(self.root)
        win.grab_set()

        tk.Label(win, text="Document Info font size:").pack(padx=10, pady=(10, 5), anchor="w")
        v = tk.IntVar(value=self.doc_info_font_size)
        sp = tk.Spinbox(win, from_=6, to=32, textvariable=v, width=4)
        sp.pack(padx=10, pady=5, anchor="w")

        def apply_and_close():
            try:
                self.doc_info_font_size = int(v.get())
            except Exception:
                self.doc_info_font_size = 10
            current_font = tkfont.Font(font=self.output_text.cget("font"))
            current_font.configure(size=self.doc_info_font_size)
            self.output_text.configure(font=current_font)
            if hasattr(self, 'current_file') and hasattr(self, 'current_text'):
                self.show_document_info(self.current_file, self.current_text)
            self.save_dashboard()
            win.destroy()

        btns = tk.Frame(win)
        btns.pack(fill="x", padx=10, pady=(10,10))
        ttk.Button(btns, text="OK", command=apply_and_close).pack(side="right")
        ttk.Button(btns, text="Cancel", command=win.destroy).pack(side="right", padx=(0,5))

    def _open_dashboard_dialog(self):
        win = tk.Toplevel(self.root)
        win.title("Dashboard")
        win.transient(self.root)
        win.grab_set()
        frm = tk.Frame(win, padx=6, pady=6)
        frm.pack(fill="both", expand=True)

        # Amplifiers
        amp = tk.LabelFrame(frm, text="Amplifiers", padx=6, pady=6, bg="#e0f7ff")
        amp.grid(row=0, column=0, sticky="nsew", padx=4, pady=4)
        tk.Label(amp, text="Factor:", bg="#e0f7ff").grid(row=0, column=0, sticky="w")
        amp_factor_var = tk.StringVar(value=str(self.amplifiers["factor"]))
        tk.Entry(amp, width=6, textvariable=amp_factor_var).grid(row=0, column=1, sticky="w", padx=(4,6))
        tk.Label(amp, text="Add word:", bg="#e0f7ff").grid(row=1, column=0, sticky="w")
        amp_word = tk.Entry(amp, width=16)
        amp_word.grid(row=1, column=1, sticky="w", padx=(4,6))
        amp_list = tk.Listbox(amp, height=8)
        amp_list.grid(row=2, column=0, columnspan=3, sticky="nsew", pady=(6,0))
        amp.grid_rowconfigure(2, weight=1)
        amp.grid_columnconfigure(2, weight=1)

        def amp_refresh():
            amp_list.delete(0, tk.END)
            # Display all CURRENTLY active amplifiers (defaults + file parts)
            for w in self.amplifiers["words"]:
                amp_list.insert(tk.END, w)

        def amp_add():
            w = amp_word.get().strip()
            if w and w not in self.amplifiers["words"]:
                self.amplifiers["words"].append(w)
                # Also add to default_amplifiers so it persists across sessions
                if w not in self.default_amplifiers:
                    self.default_amplifiers.append(w)
                amp_word.delete(0, tk.END)
                amp_refresh()

        def amp_remove():
            sel = list(amp_list.curselection())
            for i in reversed(sel):
                word_to_remove = amp_list.get(i)
                
                # Remove from active session (handling potential duplicates)
                while word_to_remove in self.amplifiers["words"]:
                    self.amplifiers["words"].remove(word_to_remove)
                    
                # Remove from persistent defaults so it stays gone
                while word_to_remove in self.default_amplifiers:
                    self.default_amplifiers.remove(word_to_remove)
                    
            amp_refresh()

        btn_row = tk.Frame(amp, bg="#e0f7ff")
        btn_row.grid(row=1, column=2, sticky="w")
        ttk.Button(btn_row, text="Add", command=amp_add).pack(side="left", padx=2)
        ttk.Button(btn_row, text="Remove", command=amp_remove).pack(side="left", padx=2)

        amp_refresh()

        # Silencers
        sil = tk.LabelFrame(frm, text="Silencers", padx=6, pady=6, bg="#ffe0e0")
        sil.grid(row=0, column=1, sticky="nsew", padx=4, pady=4)
        tk.Label(sil, text="Factor:", bg="#ffe0e0").grid(row=0, column=0, sticky="w")
        sil_factor_var = tk.StringVar(value=str(self.silencers["factor"]))
        tk.Entry(sil, width=6, textvariable=sil_factor_var).grid(row=0, column=1, sticky="w", padx=(4,6))
        tk.Label(sil, text="Add word:", bg="#ffe0e0").grid(row=1, column=0, sticky="w")
        sil_word = tk.Entry(sil, width=16)
        sil_word.grid(row=1, column=1, sticky="w", padx=(4,6))
        sil_list = tk.Listbox(sil, height=8)
        sil_list.grid(row=2, column=0, columnspan=3, sticky="nsew", pady=(6,0))
        sil.grid_rowconfigure(2, weight=1)
        sil.grid_columnconfigure(2, weight=1)

        def sil_refresh():
            sil_list.delete(0, tk.END)
            seen=set()
            for w in self.silencers["words"]:
                if isinstance(w, str) and w.lower() not in seen:
                    sil_list.insert(tk.END, w)
                    seen.add(w.lower())

        def sil_add():
            w = sil_word.get().strip()
            if w and w not in self.silencers["words"]:
                self.silencers["words"].append(w)
                sil_word.delete(0, tk.END)
                sil_refresh()

        def sil_remove():
            sel = list(sil_list.curselection())
            for i in reversed(sel):
                del self.silencers["words"][i]
            sil_refresh()

        btn_row2 = tk.Frame(sil, bg="#ffe0e0")
        btn_row2.grid(row=1, column=2, sticky="w")
        ttk.Button(btn_row2, text="Add", command=sil_add).pack(side="left", padx=2)
        ttk.Button(btn_row2, text="Remove", command=sil_remove).pack(side="left", padx=2)

        sil_refresh()

        # Length Penalty
        len_penalty = tk.LabelFrame(frm, text="Length Penalty", padx=6, pady=6, bg="#eef0f2")
        len_penalty.grid(row=1, column=0, sticky="nsew", padx=4, pady=4)
        tk.Label(len_penalty, text="Word Threshold:", bg="#eef0f2").grid(row=0, column=0, sticky="w")
        len_thresh_var = tk.StringVar(value=str(self.length_penalty.get("threshold", 3000)))
        tk.Entry(len_penalty, width=8, textvariable=len_thresh_var).grid(row=0, column=1, sticky="w", padx=(4,6))
        tk.Label(len_penalty, text="Penalty Factor:", bg="#eef0f2").grid(row=1, column=0, sticky="w")
        len_factor_var = tk.StringVar(value=str(self.length_penalty.get("factor", 0.9)))
        tk.Entry(len_penalty, width=8, textvariable=len_factor_var).grid(row=1, column=1, sticky="w", padx=(4,6))

        # Recent Files & Max Recs
        recent_files_frm = tk.LabelFrame(frm, text="General Settings", padx=6, pady=6, bg="#f0f0e0")
        recent_files_frm.grid(row=1, column=1, sticky="nsew", padx=4, pady=4)
        tk.Label(recent_files_frm, text="Timespan (hours):", bg="#f0f0e0").grid(row=0, column=0, sticky="w")
        recent_timespan_var = tk.StringVar(value=str(self.recent_timespan_hours))
        tk.Entry(recent_files_frm, width=8, textvariable=recent_timespan_var).grid(row=0, column=1, sticky="w", padx=(4,6))
        
        tk.Label(recent_files_frm, text="Max Recommendations:", bg="#f0f0e0").grid(row=1, column=0, sticky="w")
        max_recs_var = tk.StringVar(value=str(self.max_recommendations))
        tk.Entry(recent_files_frm, width=8, textvariable=max_recs_var).grid(row=1, column=1, sticky="w", padx=(4,6))

        # footer
        footer = tk.Frame(frm)
        footer.grid(row=2, column=0, columnspan=2, sticky="e", pady=(6,0))
        def apply_close():
            # persist factors
            try:
                self.amplifiers["factor"] = float(amp_factor_var.get())
            except Exception:
                pass
            try:
                self.silencers["factor"] = float(sil_factor_var.get())
            except Exception:
                pass
            try:
                self.length_penalty["threshold"] = int(len_thresh_var.get())
            except Exception:
                pass
            try:
                self.length_penalty["factor"] = float(len_factor_var.get())
            except Exception:
                pass
            try:
                self.recent_timespan_hours = int(recent_timespan_var.get())
            except Exception:
                self.recent_timespan_hours = 24 # Fallback
            try:
                self.max_recommendations = int(max_recs_var.get())
            except Exception:
                self.max_recommendations = 20 # Fallback

            self.save_dashboard()
            win.destroy()

        ttk.Button(footer, text="OK", command=apply_close).pack(side="right", padx=(6,0))
        ttk.Button(footer, text="Cancel", command=win.destroy).pack(side="right")

        frm.grid_rowconfigure(0, weight=1)
        frm.grid_columnconfigure(0, weight=1)
        frm.grid_columnconfigure(1, weight=1)

    def _open_default_dir_dialog(self):
        """Opens a dialog to view and change the default directory."""
        win = tk.Toplevel(self.root)
        win.title("Default Directory")
        win.transient(self.root)
        win.grab_set()

        frm = tk.Frame(win, padx=10, pady=10)
        frm.pack(fill="both", expand=True)

        tk.Label(frm, text="Current Default Directory:").pack(anchor="w")

        dir_var = tk.StringVar(value=self.default_directory)
        entry = tk.Entry(frm, textvariable=dir_var, width=60)
        entry.pack(fill="x", expand=True, pady=(2, 5))

        def select_new_dir():
            directory = filedialog.askdirectory(initialdir=self.default_directory or os.getcwd())
            if directory:
                dir_var.set(directory)

        tk.Button(frm, text="Browse...", command=select_new_dir).pack(anchor="w", pady=(0, 10))

        def apply_and_close():
            new_dir = dir_var.get().strip()
            normalized_default_dir = str(Path(new_dir).resolve()) if new_dir else ""
            if Path(normalized_default_dir).parent == Path(normalized_default_dir) and not normalized_default_dir.endswith(os.sep):
                normalized_default_dir += os.sep
            self.default_directory = normalized_default_dir
            self.save_dashboard()
            messagebox.showinfo("Saved", "Default directory updated.", parent=win)
            win.destroy()

        btns = tk.Frame(win)
        btns.pack(fill="x", padx=10, pady=(0, 5))
        ttk.Button(btns, text="OK", command=apply_and_close).pack(side="right")
        ttk.Button(btns, text="Cancel", command=win.destroy).pack(side="right", padx=(0,5))

    # ---------------------------------------

    def toggle_all_directories(self):
        if not self.directories:
            return
        self._clear_active_preset()
            
        # Determine whether to turn all ON or all OFF
        all_on = all(self.directory_active_status.get(d, True) for d in self.directories)
        new_status = not all_on
        
        for directory in self.directories:
            self.directory_active_status[directory] = new_status
            
        self.save_dashboard()
        self.refresh_toggle_buttons()
        
        # Update TreeView styles to reflect the new state
        for iid in self.dir_tree.get_children():
            try:
                val = self.dir_tree.item(iid, "values")[0]
                node_norm = str(Path(val).resolve())
                
                # Check if this node corresponds to a root directory
                if node_norm in self.directories:
                    base_color = self.directory_colors.get(node_norm, "#ffffff")
                    root_color = self._get_depth_color(base_color, 0)
                    fg_color = self._get_contrast_color(root_color)
                    active_tag = f"bg_{root_color}"
                    
                    self.dir_tree.tag_configure(active_tag, background=root_color, foreground=fg_color)
                    tag = active_tag if new_status else "inactive"
                    self.dir_tree.item(iid, tags=(tag,))
            except Exception:
                continue

    def apply_preset(self, p_id):
        """Applies the saved ON/OFF state of directories."""
        state = self.presets[p_id].get("state", {})
        if not state:
            messagebox.showinfo("Preset", "No state saved in this preset yet.\nRight-click the button to save your current active directories.")
            return
            
        self.active_preset = p_id             # <--- NEW
        self._update_preset_buttons_ui()      # <--- NEW
        
        for d in self.directories:
            if d in state:
                self.directory_active_status[d] = state[d]
                
        self.save_dashboard()
        self.refresh_toggle_buttons()
        
        # Update TreeView styles
        for iid in self.dir_tree.get_children():
            try:
                val = self.dir_tree.item(iid, "values")[0]
                node_norm = str(Path(val).resolve())
                
                if node_norm in self.directories:
                    is_active = self.directory_active_status.get(node_norm, True)
                    base_color = self.directory_colors.get(node_norm, "#ffffff")
                    root_color = self._get_depth_color(base_color, 0)
                    fg_color = self._get_contrast_color(root_color)
                    
                    active_tag = f"bg_{root_color}"
                    self.dir_tree.tag_configure(active_tag, background=root_color, foreground=fg_color)
                    tag = active_tag if is_active else "inactive"
                    self.dir_tree.item(iid, tags=(tag,))
            except Exception:
                continue
                
        if hasattr(self, "status_label"):
            self.status_label.config(text=f"Applied preset '{self.presets[p_id]['name']}'")

    def show_preset_menu(self, event, p_id):
        """Right-click menu to Save, Rename, or Color the preset."""
        menu = tk.Menu(self.root, tearoff=0)
        menu.add_command(label=f"Save Current State to '{self.presets[p_id]['name']}'", command=lambda: self.save_preset_state(p_id))
        menu.add_command(label="Rename Preset", command=lambda: self.rename_preset(p_id))
        menu.add_command(label="Change Active Color", command=lambda: self.change_preset_color(p_id))
        menu.tk_popup(event.x_root, event.y_root)

    def save_preset_state(self, p_id):
        self.presets[p_id]["state"] = self.directory_active_status.copy()
        self.save_dashboard()
        if hasattr(self, "status_label"):
            self.status_label.config(text=f"Saved state to preset '{self.presets[p_id]['name']}'")

    def rename_preset(self, p_id):
        from tkinter import simpledialog
        new_name = simpledialog.askstring("Rename Preset", "Enter a short name (max 6 chars):", initialvalue=self.presets[p_id]["name"])
        if new_name is not None:
            new_name = new_name.strip()
            if new_name:
                self.presets[p_id]["name"] = new_name
                self.preset_buttons[p_id].config(text=new_name)
                self.save_dashboard()

    def create_widgets(self):
        # ——— Sidebar: Cached Directories ———
        # ——— Sidebar: Cached Directories ———
        tk.Label(self.sidebar, text="Cached Directories:", font=("Helvetica", 11, "bold")) \
            .pack(anchor="w", padx=5, pady=(5,0))

        # Create a container frame for the two buttons
        sidebar_top_btn_frame = tk.Frame(self.sidebar)
        sidebar_top_btn_frame.pack(fill="x", padx=5, pady=5)

        self.view_mode = tk.StringVar(value="Cached tree")
        self.view_mode_button = tk.Button(
            sidebar_top_btn_frame,
            text=self.view_mode.get(),
            bg="#ccffcc",
            width=12,
            command=self.toggle_view_mode
        )
        self.view_mode_button.pack(side="left")

        tk.Button(
            sidebar_top_btn_frame,
            text="Reload All",
            command=self.reload_all_directories,
            bg="#ffd700"
        ).pack(side="right")

        search_frame = tk.Frame(self.sidebar)
        search_frame.pack(fill="x", padx=5, pady=(0,5))
        
        # Replace the old "Search (ayo):" label with the clickable help label
        self.search_help_label = tk.Label(search_frame, text="❓ Search help", font=("Helvetica", 8, "underline", "bold"), fg="blue", cursor="hand2")
        self.search_help_label.pack(side="left")
        self.search_help_label.bind("<Button-1>", lambda e: self.show_search_help())

        self.search_var = tk.StringVar()
        search_entry = tk.Entry(search_frame, textvariable=self.search_var)
        search_entry.pack(side="left", fill="x", expand=True, padx=(5,0))
        search_entry.bind("<Return>", self.on_search_execute)
        tk.Button(search_frame, text="Go", command=self.on_search_execute).pack(side="right", padx=(5,0))

        self.sidebar_tree_frame = ttk.Frame(self.sidebar)
        self.sidebar_tree_frame.pack(fill="both", expand=True, padx=5, pady=5)

        self.tree_vscroll = ttk.Scrollbar(self.sidebar_tree_frame, orient="vertical")
        self.tree_vscroll.pack(side="right", fill="y")

        self.dir_tree = ttk.Treeview(
            self.sidebar_tree_frame,
            show="tree",
            yscrollcommand=self.tree_vscroll.set
        )
        self.dir_tree.pack(side="left", fill="both", expand=True)
        self.tree_vscroll.config(command=self.dir_tree.yview)

        self.dir_tree.tag_configure("active",   background="#ffffcc")
        self.dir_tree.tag_configure("inactive", background="#f0f0f0", foreground="black")

        self.dir_tree.bind("<Button-3>", self._on_tree_right_click)
        self.dir_tree.bind("<<TreeviewOpen>>",   self._on_tree_expand)
        self.dir_tree.bind("<ButtonRelease-1>",  self._on_tree_single_click)
        self.dir_tree.bind("<Double-1>",       self._on_tree_double_click)

        # --- NEW TOGGLE BUTTONS FRAME ---
        toggle_header_frame = tk.Frame(self.sidebar)
        toggle_header_frame.pack(fill="x", padx=5, pady=(10,0))
        
        tk.Button(toggle_header_frame, text="Toggle Directories:", font=("Helvetica", 10, "bold"),
                  command=self.toggle_all_directories).pack(side="left")
                  
        self.preset_buttons = {}
        for p_id in ["1", "2", "3"]:
            btn = tk.Button(toggle_header_frame, text=self.presets[p_id]["name"], bg="#e0e0e0",
                            command=lambda pid=p_id: self.apply_preset(pid))
            btn.pack(side="left", padx=2)
            # Right-click to show menu (Save/Rename)
            btn.bind("<Button-3>", lambda e, pid=p_id: self.show_preset_menu(e, pid))
            self.preset_buttons[p_id] = btn
        # ---------------------------------

        self.toggle_btns_frame = ttk.Frame(self.sidebar)
        self.toggle_btns_frame.pack(fill="x", padx=5, pady=(0,5))

        self.refresh_toggle_buttons()

        self.dir_menu = tk.Menu(self.sidebar, tearoff=0)
        self.dir_menu.add_command(label="Save On/Off", command=self._toggle_save_selected_file)
        self.dir_menu.add_command(label="RUN", command=self._run_selected_file)
        self.dir_menu.add_command(label="I have thoughts", command=self._i_have_thoughts) # <--- Add this
        self.dir_menu.add_command(label="OPEN", command=self._open_selected_file_with_os)
        self.dir_menu.add_command(label="Rename", command=self._rename_selected_file)
        self.dir_menu.add_command(label="Folder", command=self._open_selected_file)
        self.dir_menu.add_command(label="Zoom", command=self._zoom_selected_dirs)
        self.dir_menu.add_command(label="Copy path", command=self._copy_path_to_clipboard)
        self.dir_menu.add_command(label="Tree", command=self._focus_on_selected_in_tree)
        self.dir_menu.add_command(label="OCR (Image)", command=self._run_ocr_on_selected)

        self.dir_tree.bind("<Button-3>", self._on_tree_right_click)
        self.dir_tree.bind("<<TreeviewOpen>>", self._on_tree_expand)
        self.dir_tree.bind("<ButtonRelease-1>", self._on_tree_single_click)
        self.dir_tree.bind("<Double-1>", self._on_tree_double_click)

        top_panel = tk.Frame(self.right_pane, bd=2, relief="raised", padx=5, pady=5)
        top_panel.pack(fill="x", pady=(5,10))
        tk.Button(top_panel, text="Add Dir", command=self.add_directory, bg="#cce5ff").pack(side="left", padx=3)
        tk.Button(top_panel, text="Select File", command=self.select_current_document, bg="#ffebcc").pack(side="left", padx=3)
        tk.Button(top_panel, text="Rerun", command=self.run_dashboard, bg="#ffcccc").pack(side="left", padx=3)

        # Options dropdown (Default Dir, Font, Dashboard)
        opt_btn = tk.Menubutton(top_panel, text="Options", relief="raised")
        opt_menu = tk.Menu(opt_btn, tearoff=0)
        opt_menu.add_command(label="Default Dir…", command=self._open_default_dir_dialog)
        opt_menu.add_command(label="Font…", command=self._open_font_dialog)
        opt_menu.add_command(label="Dashboard…", command=self._open_dashboard_dialog)
        opt_menu.add_separator()
        opt_menu.add_command(label="Year Lookback", command=self.generate_year_lookback)
        opt_btn.config(menu=opt_menu)
        opt_btn.pack(side="left", padx=(10,3))

        self.status_label = tk.Label(top_panel, text="", font=("Helvetica", 9, "italic"))
        self.status_label.pack(side="left", padx=10)

        self.toggle_rec_btn = tk.Button(top_panel, text="Hide", command=self.toggle_recommendations, bg="#d3d3d3")
        self.toggle_rec_btn.pack(side="right", padx=3)

        # --- Replaced standard frame split with an adjustable ttk.Panedwindow ---
        self.middle_paned = ttk.Panedwindow(self.right_pane, orient="horizontal")
        self.middle_paned.pack(fill="both", expand=True, pady=10)

        self.about_frame = tk.Frame(self.middle_paned, bd=1, relief="groove")
        self.middle_paned.add(self.about_frame, weight=1)

        header_frame = tk.Frame(self.about_frame)
        header_frame.pack(fill="x", padx=5, pady=5)
        tk.Label(header_frame, text="Document Info:", font=("Helvetica", 12, "bold")).pack(side="left")
        self.doc_info_title_label = tk.Label(header_frame, text="No file selected", font=("Helvetica", 10, "bold"), anchor="w")
        self.doc_info_title_label.pack(side="left", padx=(5,0), fill="x", expand=True)

        # Filepath label
        if not hasattr(self, "filepath_var"):
            self.filepath_var = tk.StringVar(value="")
        tk.Label(self.about_frame, textvariable=self.filepath_var, font=("Helvetica", 9, "bold"), anchor="w").pack(fill="x", padx=5, pady=(0, 0))

        # Published label is now placed here, below the title and above the text box
        if not hasattr(self, "published_var"):
            self.published_var = tk.StringVar(value="Published: —")
        tk.Label(self.about_frame, textvariable=self.published_var, font=("Helvetica", 9, "bold"), anchor="w").pack(fill="x", padx=5, pady=(0,5))

        text_frame = tk.Frame(self.about_frame)
        text_frame.pack(fill="both", expand=True, padx=5, pady=5)
        scroll_bar = ttk.Scrollbar(text_frame, orient="vertical")
        scroll_bar.pack(side="left", fill="y")
        self.output_text = tk.Text(text_frame, wrap="word", font=("Verdana", self.doc_info_font_size), yscrollcommand=scroll_bar.set)
        self.output_text.pack(side="left", fill="both", expand=True)
        scroll_bar.config(command=self.output_text.yview)
        scroll_bar.lift()
        self.output_text.config(state="disabled")

        self.rec_frame = tk.Frame(self.middle_paned, bd=1, relief="groove")
        tk.Label(self.rec_frame, text="Recommendations:", font=("Helvetica", 12, "bold")).pack(anchor="w", padx=5, pady=(5,0))
        
        # --- NEW OPTIONS FRAME FOR INLINE CHECKBOXES ---
        opts_frame = tk.Frame(self.rec_frame)
        opts_frame.pack(fill="x", anchor="w", padx=5, pady=(0,5))
        
        tk.Checkbutton(opts_frame, text="Sort by Filename", variable=self.sort_by_filename,
                  command=self.refresh_recommendation_buttons).pack(side="left")
        tk.Checkbutton(opts_frame, text="Amplify", variable=self.amplify_results,
                  command=self._on_amplify_toggle).pack(side="left", padx=(10, 0))
        # -----------------------------------------------

        self.shortcut_buttons_frame = tk.Frame(self.rec_frame)
        self.shortcut_buttons_frame.pack(fill="both", expand=True, padx=5, pady=5)

        # Add this line right next to your other self.dir_tree.bind commands
        self.dir_tree.bind("<<TreeviewSelect>>", self._on_tree_select)
        
        self.dir_tree.bind("<Button-3>", self._on_tree_right_click)
        self.dir_tree.bind("<<TreeviewOpen>>",   self._on_tree_expand)
        self.dir_tree.bind("<ButtonRelease-1>",  self._on_tree_single_click)
        self.dir_tree.bind("<Double-1>",       self._on_tree_double_click)

        # Start with recommendations hidden
        self.toggle_rec_btn.config(text="Show")

    def _on_closing(self):
        """Handle the window closing event."""
        self.check_and_update_last_opened_file()
        self.save_dashboard()
        self.root.destroy()

    def _rename_selected_file(self):
        """Renames the selected file in the Treeview and disk, then refreshes the directory."""
        from tkinter import simpledialog
        
        selected_item = self.dir_tree.focus()
        if not selected_item:
            return
            
        old_path = self._node_path(selected_item)
        if not os.path.exists(old_path) or not os.path.isfile(old_path):
            messagebox.showwarning("Rename", "Please select a valid file to rename.")
            return
            
        old_name = os.path.basename(old_path)
        new_name = simpledialog.askstring("Rename File", "Enter new filename (including extension):", initialvalue=old_name)
        
        if new_name and new_name != old_name:
            new_path = os.path.join(os.path.dirname(old_path), new_name)
            try:
                os.rename(old_path, new_path)
                
                # Update the active UI if it's the currently opened file
                if getattr(self, "current_document", None) == old_path:
                    self.current_document = new_path
                    self.doc_info_title_label.config(text=new_name)
                
                # Refresh the parent root directory to safely update the cache and tree
                root_dir = self._get_tree_root_path(selected_item)
                if root_dir in self.directories:
                    self.status_label.config(text=f"Renamed to {new_name}. Updating cache...")
                    self._call_build_cache(root_dir, force_rebuild=False)
                    
            except Exception as e:
                messagebox.showerror("Rename Error", f"Could not rename file:\n{e}")

    def _open_selected_file(self):
        selected_item = self.dir_tree.focus()
        if selected_item:
            file_path = self.dir_tree.item(selected_item, "values")[0]
            if os.path.isfile(file_path):
                folder_path = os.path.dirname(file_path)
            elif os.path.isdir(file_path):
                folder_path = file_path
            else:
                messagebox.showerror("Error", "Selected item is neither a file nor a valid directory.")
                return
            if folder_path:
                try:
                    open_file(folder_path)
                except Exception as e:
                    messagebox.showerror("Error", f"Could not open folder:\n{e}")

    def toggle_recommendations(self):
        if str(self.rec_frame) in self.middle_paned.panes():
            self.middle_paned.forget(self.rec_frame)
            self.toggle_rec_btn.config(text="Show")
        else:
            self.middle_paned.add(self.rec_frame, weight=1)
            self.toggle_rec_btn.config(text="Hide")

    def ensure_recommendations_visible(self):
        if str(self.rec_frame) not in self.middle_paned.panes():
            self.middle_paned.add(self.rec_frame, weight=1)
            self.toggle_rec_btn.config(text="Hide")

    def _zoom_selected_dirs(self):
        apply_zoom_view(self)

    def on_search_execute(self, event=None):
        query = self.search_var.get().strip()
        search.execute_search(
            query=query,
            cache=self.cache,
            directory_active_status=self.directory_active_status,
            dir_tree=self.dir_tree,
            on_view_mode_change_callback=lambda *_: cached_saved_recent_button.on_view_mode_change(self),
            app=self  # <--- ADD THIS LINE
        )

    def _run_selected_file(self):
        sel = self.dir_tree.selection()
        if not sel:
            return
        iid = sel[0]
        path = self._node_path(iid)
        if not os.path.isfile(path):
            return
        self.dig_document(path)

    def toggle_view_mode(self):
        """Handles view mode toggling by calling the external module."""
        current_mode = self.view_mode.get()
        
        # Get the next mode and color from the module
        new_mode, color = cached_saved_recent_button.get_next_view_mode(current_mode)
        
        # Update the UI state
        self.view_mode.set(new_mode)
        self.view_mode_button.config(text=new_mode, bg=color)
        
        # Trigger the tree rebuild using the function from the module
        cached_saved_recent_button.on_view_mode_change(self)

    def _on_tree_single_click(self, event):
        self.check_and_update_last_opened_file()
        iid = self.dir_tree.identify_row(event.y)
        if not iid:
            return
        path = self._node_path(iid)
        if os.path.isfile(path):
            self.current_document = path
            self.doc_info_title_label.config(text=os.path.basename(path))
            
            # Just extract existing cache/text, no OCR on click
            text = cache_generator.extract_text(path)
            self.show_document_info(path, text)

    def _on_tree_double_click(self, event):
        iid = self.dir_tree.identify_row(event.y)
        if not iid:
            return
        path = self._node_path(iid)
        if os.path.isfile(path):
            self._open_file_and_monitor(path)
        elif os.path.isdir(path):       # <--- ADD THIS
            self.find_in_tree(path)     # <--- ADD THIS

    def _on_tree_expand(self, event):
        sel_list = self.dir_tree.selection()
        if not sel_list:
            return
        self._expand_and_populate(sel_list[0])

    def _on_tree_select(self, event):
        self.check_and_update_last_opened_file()
        sel = self.dir_tree.selection()
        if not sel:
            return
            
        path = self._node_path(sel[0])
        if os.path.isfile(path):
            self.current_document = path
            self.doc_info_title_label.config(text=os.path.basename(path))
            
            # Extract existing cache/text to display in the preview panel
            import cache_generator # Ensure this is accessible if not already imported at the top
            text = cache_generator.extract_text(path)
            self.show_document_info(path, text)

    def _node_path(self, item):
        item_values = self.dir_tree.item(item, "values")
        if item_values and item_values[0]:
            return str(Path(item_values[0]).resolve())
        return ""

    def reload_all_directories(self):
        import threading
        import time
        import recommender
        
        self.status_label.config(text="Reloading all directories in background...")
        directories_to_reload = list(self.directories.keys())

        def reload_task():
            last_update = [0]
            
            def progress_callback(message):
                now = time.time()
                # Throttle UI updates to a maximum of 10 times per second
                if now - last_update[0] > 0.1:
                    last_update[0] = now
                    # Safely pass the specific message to the Tkinter main loop
                    self.root.after(0, lambda msg=message: self.status_label.config(text=msg))

            # Process sequentially to prevent thread explosion
            for directory in directories_to_reload:
                updated_cache_data = cache_generator.build_cache(
                    directory=directory,
                    cache=self.cache,
                    force_rebuild=False,
                    progress_callback=progress_callback
                )
                
                # Update the UI for this specific directory once it finishes
                def finalize_dir(dir_name=directory, data=updated_cache_data):
                    if data:
                        self.cache[dir_name] = data
                        self.directories[dir_name] = data["timestamp"]
                        self.update_directory_frame(dir_name)

                self.root.after(0, finalize_dir)

            # Finalize the entire process once the loop completes
            def finalize_all():
                self.save_cache()
                self.save_dashboard()
                self.status_label.config(text="Rebuilding Recommender Model...")
                
                # Offload the recommender model rebuild to a separate thread
                def rebuild_model_task():
                    try:
                        recommender.update_model(self.cache)
                        self.root.after(0, lambda: self.status_label.config(text="All directories reloaded and model rebuilt!"))
                        
                        # Refresh recommendations if a document is currently open
                        if getattr(self, 'current_document', None):
                            self.root.after(0, lambda: self.recommend_similar_files(self.current_document))
                    except Exception as e:
                        print(f"Model update failed: {e}")
                        self.root.after(0, lambda: self.status_label.config(text="Model rebuild failed!"))

                threading.Thread(target=rebuild_model_task, daemon=True).start()

            self.root.after(0, finalize_all)

        # Launch the single sequential background task
        threading.Thread(target=reload_task, daemon=True).start()

    # (amp/sil list helpers kept for dialogs reuse)
    def refresh_amp_list(self):
        pass  # now handled inside the dialog

    def refresh_sil_list(self):
        pass  # now handled inside the dialog

    def remove_amplifier_word(self, index):
        if 0 <= index < len(self.amplifiers["words"]):
            del self.amplifiers["words"][index]
            self.save_dashboard()

    def remove_silencer_word(self, index):
        if 0 <= index < len(self.silencers["words"]):
            del self.silencers["words"][index]
            self.save_dashboard()

    def run_dashboard(self):
        self.ensure_recommendations_visible()
        self.save_dashboard()
        if self.current_document:
            text = cache_generator.extract_text(self.current_document)
            self.show_document_info(self.current_document, text)
            self.clear_shortcut_buttons()
            self.recommend_similar_files(self.current_document, current_text=text)

    def save_dashboard(self):
        # Only save the clean defaults to the txt file
        save_amplifiers_to_txt(getattr(self, 'default_amplifiers', self.amplifiers["words"]))
        
        data = {
            "amplifiers": {"words": self.amplifiers["words"], "factor": self.amplifiers["factor"]},
            "silencers": self.silencers,
            "length_penalty": self.length_penalty,
            "recent_timespan_hours": self.recent_timespan_hours,
            "max_recommendations": self.max_recommendations,
            "default_directory": self.default_directory,
            "directory_active_status": self.directory_active_status,
            "directories": self.directories,
            "directory_colors": self.directory_colors,
            "presets": self.presets,
            "saved_files": list(self.saved_files),
            "sidebar_width": self.sidebar_width,
            "doc_info_font_size": self.doc_info_font_size
        }
        with open(DASHBOARD_FILE, "w") as f:
            json.dump(data, f)

# --- NEW: Automatically export saved files to a TXT file ---
        try:
            export_path = Path(__file__).parent / "saved_files.txt"
            with open(export_path, "w", encoding="utf-8") as f:
                # Sort for a clean, predictable list
                for file_path in sorted(list(self.saved_files)):
                    f.write(f"{file_path}\n")
        except Exception as e:
            print(f"[ERROR] Could not export saved_files.txt: {e}")

    def load_dashboard(self):
        if DASHBOARD_FILE.exists():
            try:
                with open(DASHBOARD_FILE, "r") as f:
                    data = json.load(f)
                self.presets = data.get("presets", {
                    "1": {"name": "1", "state": {}},
                    "2": {"name": "2", "state": {}},
                    "3": {"name": "3", "state": {}}
                })
                # Ensure existing saves get a default color assigned
                for p in self.presets.values():
                    if "color" not in p:
                        p["color"] = "#e0e0e0"

                saved_amp = data.get("amplifiers", {})
                self.default_amplifiers = load_amplifiers_from_txt()
                self.amplifiers = {
                    "words": self.default_amplifiers.copy(),
                    "factor": saved_amp.get("factor", 2.0)
                }
                self.silencers = data.get("silencers", {"words": [], "factor": 1.0})
                self.length_penalty = data.get("length_penalty", {"threshold": 3000, "factor": 0.9})
                self.recent_timespan_hours = data.get("recent_timespan_hours", 24)
                self.max_recommendations = data.get("max_recommendations", 20)
                
                # Coerce/merge legacy structures
                def _coerce_silencers(s):
                    try:
                        words = s.get("words", [])
                        if not isinstance(words, list):
                            words = list(words) if words else []
                        # Merge legacy keys
                        extra = []
                        for k in ("silenced_words", "filenames", "silencers_filenames"):
                            v = data.get(k, [])
                            if isinstance(v, list):
                                extra.extend(v)
                        allw = []
                        seen = set()
                        for w in list(words) + extra:
                            if not isinstance(w, str):
                                continue
                            wl = w.strip()
                            if not wl:
                                continue
                            if wl.lower() in seen:
                                continue
                            seen.add(wl.lower())
                            allw.append(wl)
                        return {"words": allw, "factor": float(s.get("factor", 1.0))}
                    except Exception:
                        return {"words": [], "factor": 1.0}
        
                self.silencers = _coerce_silencers(self.silencers)
                # Backward-compat: merge legacy keys
                try:
                    legacy = data.get("silenced_words", [])
                    if legacy:
                        self.silencers["words"].extend([w for w in legacy if w not in self.silencers["words"]])
                    fnames = data.get("filenames", []) or data.get("silencers_filenames", [])
                    if fnames:
                        self.silencers["words"].extend([w for w in fnames if w not in self.silencers["words"]])
                except Exception:
                    pass
                self.saved_files = set(data.get("saved_files", []))

                default_dir = data.get("default_directory", "")
                if default_dir:
                    norm = str(Path(default_dir).resolve())
                    if Path(norm).parent == Path(norm) and not norm.endswith(os.sep):
                        norm += os.sep
                    self.default_directory = norm
                else:
                    self.default_directory = ""

                normalized_active_status = {}
                for k, v in data.get("directory_active_status", {}).items():
                    nk = str(Path(k).resolve())
                    if Path(nk).parent == Path(nk) and not nk.endswith(os.sep):
                        nk += os.sep
                    normalized_active_status[nk] = v
                self.directory_active_status = normalized_active_status

                normalized_directories = {}
                for k, v in data.get("directories", {}).items():
                    nk = str(Path(k).resolve())
                    if Path(nk).parent == Path(nk) and not nk.endswith(os.sep):
                        nk += os.sep
                    normalized_directories[nk] = v
                self.directories = normalized_directories

                normalized_colors = {}
                for k, v in data.get("directory_colors", {}).items():
                    nk = str(Path(k).resolve())
                    if Path(nk).parent == Path(nk) and not nk.endswith(os.sep):
                        nk += os.sep
                    normalized_colors[nk] = v
                self.directory_colors = normalized_colors

                # Ensure all directories have a color assigned
                for d in self.directories:
                    if d not in self.directory_colors:
                        self.directory_colors[d] = self._generate_random_color()

                self.sidebar_width = data.get("sidebar_width", self.sidebar_width)
                self.doc_info_font_size = data.get("doc_info_font_size", 10)

            except Exception as e:
                print(f"[ERROR] Loading dashboard: {e}")
                self._apply_first_run_defaults()
                self.save_dashboard()
        else:
            self._apply_first_run_defaults()
            self.save_dashboard()

    def _apply_first_run_defaults(self):
        self.default_amplifiers = load_amplifiers_from_txt()
        self.amplifiers = {"words": self.default_amplifiers.copy(), "factor": 2.0}
        self.silencers = {"words": [], "factor": 1.0}
        self.length_penalty = {"threshold": 3000, "factor": 0.9}
        self.recent_timespan_hours = 24
        self.max_recommendations = 20
        self.saved_files = set()
        self.default_directory = ""
        self.directory_active_status = {}
        self.directories = {}
        self.sidebar_width = DEFAULT_SIDEBAR_WIDTH
        self.doc_info_font_size = 10

    def _get_directory_size(self, path):
        """Calculates the total size of a directory in bytes (cached for performance)."""
        # Initialize a cache dictionary if it doesn't exist yet
        if not hasattr(self, '_dir_size_cache'):
            self._dir_size_cache = {}
            
        # Return the cached size instantly if we've already calculated it
        if path in self._dir_size_cache:
            return self._dir_size_cache[path]

        total = 0
        try:
            for dirpath, _, filenames in os.walk(path):
                for f in filenames:
                    fp = os.path.join(dirpath, f)
                    if not os.path.islink(fp):
                        try:
                            total += os.path.getsize(fp)
                        except OSError:
                            pass
        except Exception:
            pass
            
        # Store the calculated size in the cache for next time
        self._dir_size_cache[path] = total
        return total

    def _sort_treeview_roots(self):
        """Sorts the root nodes of the directory tree by their size on disk."""
        roots = self.dir_tree.get_children("")
        root_sizes = []
        for iid in roots:
            path = self.dir_tree.item(iid, "values")[0]
            size = self._get_directory_size(path)
            root_sizes.append((iid, size))

        # Sort ascending (smallest first)
        root_sizes.sort(key=lambda x: x[1])

        # Re-arrange the items by moving them to their new sorted index
        for index, (iid, _) in enumerate(root_sizes):
            self.dir_tree.move(iid, "", index)

    def add_directory(self):
        directory = filedialog.askdirectory()
        if directory:
            normalized_directory = str(Path(directory).resolve())
            if Path(normalized_directory).parent == Path(normalized_directory) and not normalized_directory.endswith(os.sep):
                normalized_directory += os.sep

            if normalized_directory not in self.directories:
                self.directories[normalized_directory] = datetime.now().isoformat()
                self.directory_colors[normalized_directory] = self._generate_random_color() # <--- NEW
                self.directory_active_status[normalized_directory] = self.directory_active_status.get(normalized_directory, True)
                
                # Build tree and cache
                self.build_directory_tree(normalized_directory)
                self._sort_treeview_roots()
                self._call_build_cache(normalized_directory)
                
                # CRITICAL: Create the toggle button for the new directory
                self.refresh_toggle_buttons() 
                
                self.save_dashboard()

    def build_directory_tree(self, directory):
        for iid in self.dir_tree.get_children():
            if str(Path(self.dir_tree.item(iid, "values")[0]).resolve()) == str(Path(directory).resolve()):
                self.dir_tree.delete(iid)

        display_text = os.path.basename(directory)
        if not display_text and Path(directory).is_dir() and Path(directory).drive:
            display_text = directory
        elif not display_text:
            display_text = directory

        iid = self.dir_tree.insert("", "end", text=display_text, values=(directory,))
        self.dir_tree.insert(iid, "end", text="…")
        
        # Apply Base Color & Contrast
        base_color = self.directory_colors.get(directory, "#ffffff")
        root_color = self._get_depth_color(base_color, 0)
        fg_color = self._get_contrast_color(root_color)
        
        tag_name = f"bg_{root_color}"
        self.dir_tree.tag_configure(tag_name, background=root_color, foreground=fg_color)
        
        tag = tag_name if self.directory_active_status.get(directory, True) else "inactive"
        self.dir_tree.item(iid, tags=(tag,))

    def _expand_and_populate(self, iid):
        normalized_path = self._node_path(iid)
        self.dir_tree.delete(*self.dir_tree.get_children(iid))
        try:
            raw_entries = os.listdir(normalized_path)
            folders = []
            files = []
            
            # 1. Separate folders and files
            for name in raw_entries:
                if os.path.isdir(os.path.join(normalized_path, name)):
                    folders.append(name)
                else:
                    files.append(name)
            
            # 2. Sort folders alphabetically (A to Z)
            folders.sort(key=lambda s: s.lower())
            
            # 3. Sort files chronologically using your custom extraction logic (Newest to Oldest)
            import date_extraction
            
            # extract_year_key handles the parsed date and puts non-dated files at the bottom
            files.sort(key=date_extraction.extract_year_key, reverse=True)
            
            # 4. Combine them back together
            entries = folders + files
            
        except Exception:
            entries = []
            
        # Figure out the cascading color & contrast for this specific depth
        root_path = self._get_tree_root_path(iid)
        root_dir = next((d for d in self.directories if root_path.startswith(d)), None)
        base_color = self.directory_colors.get(root_dir, "#ffffff")
        child_depth = self._get_tree_depth(iid) + 1
        child_color = self._get_depth_color(base_color, child_depth)
        fg_color = self._get_contrast_color(child_color)
        
        tag_name = f"bg_{child_color}"
        self.dir_tree.tag_configure(tag_name, background=child_color, foreground=fg_color)
        
        for name in entries:
            child_normalized_path = os.path.join(normalized_path, name)
            try:
                if os.path.isdir(child_normalized_path):
                    cid = self.dir_tree.insert(iid, "end", text=name, values=(child_normalized_path,))
                    self.dir_tree.insert(cid, "end", text="…")
                    self.dir_tree.item(cid, tags=(tag_name,))
                else:
                    self.dir_tree.insert(iid, "end", text=name, values=(child_normalized_path,), tags=(tag_name,))
            except (PermissionError, FileNotFoundError):
                pass

    def _on_sidebar_toggle(self, dir_path):
        self.toggle_directory(dir_path)
        try:
            target_status = self.directory_active_status.get(dir_path, True)
            target_norm = str(Path(dir_path).resolve())
            
            base_color = self.directory_colors.get(dir_path, "#ffffff")
            root_color = self._get_depth_color(base_color, 0)
            fg_color = self._get_contrast_color(root_color)
            
            active_tag = f"bg_{root_color}"
            self.dir_tree.tag_configure(active_tag, background=root_color, foreground=fg_color)
            
            for iid in self.dir_tree.get_children():
                try:
                    val = self.dir_tree.item(iid, "values")[0]
                    node_norm = str(Path(val).resolve())
                    if node_norm == target_norm:
                        tag = active_tag if target_status else "inactive"
                        self.dir_tree.item(iid, tags=(tag,))
                except Exception:
                    continue
        except Exception as e:
            print(f"[WARN] TreeView update failed during toggle: {e}")

    def _on_tree_right_click(self, event):
        iid = self.dir_tree.identify_row(event.y)
        if not iid:
            return
        self.dir_tree.selection_set(iid)
        self.dir_menu.tk_popup(event.x_root, event.y_root)

    def _reload_selected_dir(self):
        d = self._node_path(self.dir_tree.selection()[0])
        self._call_build_cache(d, force_rebuild=False)

    def _erase_selected_dir(self):
        d = self._node_path(self.dir_tree.selection()[0])
        self.erase_directory(d)

    def _toggle_save_selected_file(self):
        iid = self.dir_tree.selection()[0]
        path = self._node_path(iid)
        if path in self.saved_files:
            self.saved_files.remove(path)
        else:
            self.saved_files.add(path)
        self.save_dashboard()
        if self.view_mode.get() == "Saved":
            self.dir_tree.delete(iid)

    def _copy_path_to_clipboard(self):
        iid = self.dir_tree.selection()
        if not iid: return
        path = self._node_path(iid[0])
        self.root.clipboard_clear()
        self.root.clipboard_append(path)
        #messagebox.showinfo("Copied", f"Path copied to clipboard:\n{path}")
        
    def _open_selected_file_with_os(self):
        """Opens the selected file using the OS default application."""
        selected_item = self.dir_tree.focus()
        if not selected_item:
            return
        path = self._node_path(selected_item)
        if os.path.isfile(path):
            self._open_file_and_monitor(path)

    def _focus_on_selected_in_tree(self):
        """Clears any search filter, switches to the 'Cached tree' view, and focuses on the selected file."""
        selected_item = self.dir_tree.focus()
        if not selected_item:
            return
        path = self._node_path(selected_item)
        if not path or not os.path.exists(path):
            return

        # Clear the search bar to ignore any active search filter
        self.search_var.set("")

        # Ensure the view is 'Cached tree' and trigger a full rebuild of the tree
        if self.view_mode.get() != "Cached tree":
            self.toggle_view_mode()  # This switches the mode and triggers the rebuild
        else:
            # If already in Cached tree view, manually trigger the rebuild to clear search results
            cached_saved_recent_button.on_view_mode_change(self)

        # Use 'after' to give the UI time to redraw the full tree before focusing on the path
        self.root.after(50, lambda: self.open_tree_to_path(path))
    
    def _run_ocr_on_selected(self):
        """Triggers the OCR extraction on the selected image file."""
        selected_item = self.dir_tree.focus()
        if not selected_item:
            return
            
        path = self._node_path(selected_item)
        if not os.path.isfile(path):
            return

        ext = os.path.splitext(path)[1].lower()
        if ext not in cache_generator.IMAGE_EXTENSIONS:
            messagebox.showwarning("OCR", f"Selected file is not a supported image.\nSupported: {', '.join(cache_generator.IMAGE_EXTENSIONS)}")
            return

        self.status_label.config(text=f"Running OCR on {os.path.basename(path)}...")
        self.root.update_idletasks()

        # Run in a background thread to prevent UI freezing
        import threading
        import ocr_extractor

        def ocr_task():
            success, result_text = ocr_extractor.process_single_image(path)
            self.root.after(0, lambda: self._on_ocr_complete(success, path, result_text))

        threading.Thread(target=ocr_task, daemon=True).start()

    def _on_ocr_complete(self, success, path, result_text):
        """Callback for when the threaded OCR task finishes."""
        if success:
            self.status_label.config(text=f"OCR Complete: {os.path.basename(path)}")
            # Force a cache update for this specific file so the new metadata is instantly loaded
            self.update_cache_for_single_file(path)
            
            # If the user currently has this document open, refresh the display
            if getattr(self, 'current_document', None) == path:
                self.show_document_info(path, result_text)
            
            #try:
            #    _play_chime()
            #except Exception:
            #    pass
        else:
            self.status_label.config(text="OCR Failed.")
            messagebox.showerror("OCR Error", result_text)

    def update_directory_frame(self, directory):
        frame = self.directory_frames.get(directory)
        if frame:
            for widget in frame.winfo_children():
                if isinstance(widget, tk.Label) and widget.cget("text").startswith("Loaded:"):
                    ts_str = self.directories.get(directory, None)
                    new_ts = self.format_timestamp(ts_str) if ts_str else "Never"
                    widget.config(text=f"Loaded: {new_ts}")
                    break

    def erase_directory(self, directory):
        if directory in self.directories:
            del self.directories[directory]
        if directory in self.directory_colors: # <--- NEW
            del self.directory_colors[directory]
        if directory in self.cache:
            del self.cache[directory]
        self.save_cache()
        if directory in self.directory_active_status:
            del self.directory_active_status[directory]

        for iid in self.dir_tree.get_children():
            if str(Path(self.dir_tree.item(iid, "values")[0]).resolve()) == str(Path(directory).resolve()):
                self.dir_tree.delete(iid)
                break

        self.refresh_toggle_buttons()
        self.save_dashboard()
        print(f"[LOG] Erased directory: {directory}")

    def _open_file_and_monitor(self, path):
        """Opens a file. (File monitoring has been disabled)."""
        try:
            open_file(path)
            self.status_label.config(text=f"Opened: {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Error", f"Could not open file:\n{e}")

    def check_and_update_last_opened_file(self):
        """Disabled: No longer tracking file modifications upon opening."""
        pass
    
    def update_cache_for_single_file(self, file_path):
        """Updates the cache in-memory and on-disk for a single modified file."""
        target_path_str = str(Path(file_path).resolve())
        parent_dir = None
        for d in sorted(self.cache.keys(), key=len, reverse=True):
            if target_path_str.startswith(d):
                parent_dir = d
                break

        if not parent_dir:
            print(f"[WARN] Cannot update cache for {target_path_str}; parent directory not cached.")
            return

        try:
            new_text = cache_generator.extract_text(target_path_str)
            new_mtime = os.path.getmtime(target_path_str)
        except Exception as e:
            print(f"[ERROR] Failed to read updated file {target_path_str}: {e}")
            return

        try:
            file_list = self.cache[parent_dir].get('files', [])
            idx = file_list.index(target_path_str)
            self.cache[parent_dir]['texts'][idx] = new_text
            self.cache[parent_dir]['mtimes'][idx] = new_mtime
            self.cache[parent_dir]['timestamp'] = datetime.now().isoformat()
            self.save_cache()
            self.status_label.config(text=f"Cache updated for {os.path.basename(target_path_str)}")
            print(f"[LOG] Cache successfully updated for {os.path.basename(target_path_str)}.")
        except ValueError:
            print(f"[INFO] {target_path_str} is a new file. Reloading parent directory to update cache.")
            self.status_label.config(text=f"New file detected. Reloading {os.path.basename(parent_dir)}...")
            self._call_build_cache(parent_dir, force_rebuild=False)

    def _call_build_cache(self, directory, force_rebuild=False):
        """Wrapper to call the external cache generator and update state without freezing."""
        import threading
        import time

        def cache_task():
            last_update = [0]
            
            def progress_callback(message):
                now = time.time()
                # Throttle UI updates to a maximum of 10 times per second
                if now - last_update[0] > 0.1:
                    last_update[0] = now
                    self.root.after(0, lambda msg=message: self.status_label.config(text=msg))

            # Run the heavy cache generation in the background
            updated_cache_data = cache_generator.build_cache(
                directory=directory,
                cache=self.cache,
                force_rebuild=force_rebuild,
                progress_callback=progress_callback
            )
            
            def finalize():
                # Apply the newly generated cache data back on the main thread
                if updated_cache_data:
                    self.cache[directory] = updated_cache_data
                    self.directories[directory] = self.cache[directory]["timestamp"]
                    self.save_cache()
                    self.save_dashboard()
                    self.update_directory_frame(directory)
                    self.status_label.config(text=f"Cache rebuilt: {os.path.basename(directory)}")

            # Send the finalization step back to the main UI thread
            self.root.after(0, finalize)

        # Fire off the background thread!
        threading.Thread(target=cache_task, daemon=True).start()

    def save_cache(self):
        with open(CACHE_FILE, "w") as f:
            json.dump(self.cache, f)

    def load_cache(self):
        if CACHE_FILE.exists():
            try:
                with open(CACHE_FILE, "r") as f:
                    raw_cache = json.load(f)

                normalized_cache = {}
                for k, v in raw_cache.items():
                    normalized_k = str(Path(k).resolve())
                    if Path(normalized_k).parent == Path(normalized_k) and not normalized_k.endswith(os.sep):
                        normalized_k += os.sep
                    normalized_cache[normalized_k] = v
                self.cache = normalized_cache

                for directory_key_raw, data in raw_cache.items():
                    directory = str(Path(directory_key_raw).resolve())
                    if Path(directory).parent == Path(directory) and not directory.endswith(os.sep):
                        directory += os.sep
                    if directory not in self.directories:
                        ts = data.get("timestamp", datetime.now().isoformat())
                        self.directories[directory] = ts
                        self.directory_active_status[directory] = True

                self.save_dashboard()
            except Exception as e:
                print(f"[ERROR] Loading cache: {e}")
                return

        for directory in sorted(self.directories):
            if directory in self.cache:
                ts = self.cache[directory].get("timestamp")
                if ts:
                    self.directories[directory] = ts
            self.build_directory_tree(directory)
        self._sort_treeview_roots()

        print("[LOG] Cache loaded for all known directories.")

    def select_current_document(self):
        self.check_and_update_last_opened_file()
        filetypes_str = " ".join(f"*{ext}" for ext in SUPPORTED_EXTENSIONS)
        file = filedialog.askopenfilename(
            initialdir=self.default_directory or os.getcwd(),
            filetypes=[("Supported", filetypes_str)]
        )
        if not file:
            return
        normalized_file = str(Path(file).resolve())
        self._open_file_and_monitor(normalized_file)
        self.current_document = normalized_file
        self.doc_info_title_label.config(text=os.path.basename(normalized_file))
        self.dig_document(normalized_file)
        #_play_chime()

    def display_file(self, file):
        base = os.path.splitext(os.path.basename(file))[0]
        first, rest = (base.split(" ", 1) + [""])[:2]
        words = [w for w in base.split() if w.lower() not in COMMON_WORDS]
        parts = words + ([rest] if rest else [])
        self.rest_part = rest

        if hasattr(self, "last_selected_parts"):
            for w in self.last_selected_parts:
                if w in self.amplifiers["words"]:
                    self.amplifiers["words"].remove(w)

        import datetime as _dt
        mtime = _dt.datetime.fromtimestamp(os.path.getmtime(file))
        now   = _dt.datetime.now()
        age_days = (now - mtime).days
        if age_days == 0:
            age_str = "Today"
        elif age_days == 1:
            age_str = "Yesterday"
        else:
            age_str = f"{age_days} days ago"

        ext = os.path.splitext(file)[1].lower()
        if ext == ".html":
            highlights = documentinfo.retrieve_html_highlights(file)
        elif ext == ".docx":
            highlights = documentinfo.retrieve_docx_highlights(file)
        elif ext == ".pdf":
            highlights = documentinfo.retrieve_pdf_highlights(file)
        else:
            highlights = []

        hl_intro = f"\n\nFound {len(highlights)} highlighted passage(s):\n\n" if highlights else "\n\nNo highlights found.\n\n"

        self.rest_part = rest

        full_text = (
            f"File: {os.path.basename(file)}\n"
            f"Last modified: {mtime.strftime('%Y-%m-%d %H:%M')}"
            f"  ({age_str})\n"
            + hl_intro
            + rest
        )

        self.output_text.config(state="normal")
        self.output_text.delete("1.0", tk.END)
        self.output_text.insert(tk.END, full_text)

        for color_hex, hl_text, style_dict in highlights:
            tag_name = f"hl_{color_hex}_{abs(hash(hl_text))}"
            if tag_name not in self.output_text.tag_names():
                self.output_text.tag_configure(tag_name, **style_dict)
            self.output_text.insert(tk.END, hl_text + "\n\n", tag_name)

        self.output_text.config(state="disabled")

        base = os.path.splitext(os.path.basename(file))[0]
        first, rest = (base.split(" ", 1) + [""])[:2]
        words = [w for w in base.split() if w.lower() not in COMMON_WORDS]
        parts = words + ([rest] if rest else [])
        self.rest_part = rest

        # Reset active words strictly to Defaults + Current File Parts
        self.amplifiers["words"] = self.default_amplifiers.copy() + parts
        self.search_var.set(" ".join(words))
        self.save_dashboard()
        self.current_document = file

    def show_document_info(self, file_path, file_text):
        self.current_file = file_path
        self.current_text = file_text
        self.filepath_var.set(os.path.dirname(file_path))

        documentinfo.display_doc_info(
            output_text_widget=self.output_text,
            published_var=self.published_var,
            file_path=file_path,
            file_text=file_text,
            search_query=self.search_var.get(),
            font_size=self.doc_info_font_size
        )
        # Trigger the coloring and flashing
        self._highlight_and_flash_search_terms()

    def show_search_help(self):
        help_text = """
Search Features:

1. Standard Search
   Type words normally. Finds files containing ALL the words in their name, folder path, or text content.
   Example: climate report

2. Phrase Search " "
   Use double quotes to search for an exact phrase.
   Example: "climate change"

3. Path/Name Only ( )
   Wrap your search in parentheses to ONLY search within file names and folder names (ignores file text).
   Example: (diário)
   
4. Wildcards * and ?
   Use * at the end of a word to match any ending (e.g., auto* matches automatic, automobile).
   Use ? at the end of a word to match exactly one extra character.

5. Year Sort 'ayo'
   Add 'ayo' to the end of your search (or type 'ayo' alone) to sort results by the year found in the file names (descending).
   Example: report ayo
        """
        messagebox.showinfo("Search Help", help_text.strip())

    def refresh_recommendation_buttons(self):
        # 1. Pick the correct list based on the toggle state
        is_amplified = self.amplify_results.get()
        active_list = getattr(self, "recommended_docs_amplified", []) if is_amplified else getattr(self, "recommended_docs_base", [])
        
        self.clear_shortcut_buttons()
        
        # 2. Sort and display
        if active_list:
            if self.sort_by_filename.get():
                ranked = sorted(active_list, key=lambda x: os.path.basename(x[1]).lower(), reverse=True)
            else:
                ranked = sorted(active_list, key=lambda x: x[0], reverse=True)
                
            # Limit to configured maximum recommendations
            ranked = ranked[:self.max_recommendations]
            
            self._display_recommendation_buttons(ranked)

    def _on_amplify_toggle(self):
        """Instantly swaps the UI between amplified and base results."""
        self.refresh_recommendation_buttons()

    def clear_shortcut_buttons(self):
        for widget in self.shortcut_buttons_frame.winfo_children():
            widget.destroy()

    def dig_document(self, path):
        self.check_and_update_last_opened_file()
        self.ensure_recommendations_visible()

        base = os.path.splitext(os.path.basename(path))[0]
        first, rest = (base.split(" ", 1) + [""])[:2]
        words = [w for w in base.split() if w.lower() not in COMMON_WORDS]
        parts = words + ([rest] if rest else [])
        self.rest_part = rest

        # Reset active words strictly to Defaults + Current File Parts
        self.amplifiers["words"] = self.default_amplifiers.copy() + parts

        self.search_var.set(" ".join(words))
        self.save_dashboard()

        self.doc_info_title_label.config(text=os.path.basename(path))
        print(f"[LOG] DIG initiated for: {path}")
        
        # Grab existing text
        text = cache_generator.extract_text(path)
        
        # Check if it's an image that needs OCR
        ext = os.path.splitext(path)[1].lower()
        if not text.strip() and ext in cache_generator.IMAGE_EXTENSIONS:
            self.status_label.config(text=f"Running background OCR for {os.path.basename(path)}...")
            self.root.update_idletasks()
            
            import threading
            import ocr_extractor
            
            def ocr_and_dig_task():
                success, result_text = ocr_extractor.process_single_image(path)
                self.root.after(0, lambda: self._finalize_dig_after_ocr(success, path, result_text))
                
            threading.Thread(target=ocr_and_dig_task, daemon=True).start()
        else:
            # Not an image, or already has text. Proceed normally.
            self.show_document_info(path, text)
            self.recommend_similar_files(path, current_text=text)
        
        #try:
        #    _play_chime()
        #except Exception:
        #    pass

    def _finalize_dig_after_ocr(self, success, path, result_text):
        """Callback to finish digging a document once background OCR is done."""
        if success:
            self.status_label.config(text="Auto-OCR Complete.")
            self.update_cache_for_single_file(path)
            text_to_use = result_text
        else:
            self.status_label.config(text="Auto-OCR found no text.")
            text_to_use = ""

        self.show_document_info(path, text_to_use)
        self.recommend_similar_files(path, current_text=text_to_use)

    def add_filename_to_silencers(self, target_path):
        """Add the selected file's base filename (no extension) to the silenced words list."""
        base = os.path.splitext(os.path.basename(target_path))[0]
        if not base:
            return
        existing_lower = {w.lower() for w in self.silencers["words"]}
        if base.lower() not in existing_lower:
            self.silencers["words"].append(base)
            self.save_dashboard()
        if hasattr(self, "status_label"):
            self.status_label.config(text=f"Silenced: {base}")

    def update_about_document(self, event, file_path):
        self.check_and_update_last_opened_file()
        if self.current_active_label is not None and self.current_active_label.winfo_exists():
            self.current_active_label.config(font=self._rec_font_normal)
        else:
            self.current_active_label = None
        event.widget.config(font=self._rec_font_bold)
        self.current_active_label = event.widget
        self.current_active_file = file_path
        text = cache_generator.extract_text(file_path)
        self.show_document_info(file_path, text)
        print(f"[LOG] Updated About panel for: {file_path}")

    def _display_recommendation_buttons(self, ranked):
        self.clear_shortcut_buttons()
        if ranked:
            for score, path in ranked:
                # Determine the parent directory to grab its color
                parent_dir = next((d for d in self.directories if path.startswith(d)), None)
                base_color = self.directory_colors.get(parent_dir, "#ffffff")
                
                # Calculate depth relative to the parent directory
                if parent_dir:
                    try:
                        rel = Path(path).relative_to(parent_dir)
                        depth = len(rel.parts) - 1
                    except ValueError:
                        depth = 0
                else:
                    depth = 0
                    
                # Use the cascading color function and contrast text!
                row_bg = self._get_depth_color(base_color, depth)
                    
                txt_color = self._get_contrast_color(row_bg)

                frame = tk.Frame(self.shortcut_buttons_frame, pady=2, bg=row_bg)
                frame.pack(fill="x", anchor="w", padx=5)

                btn_frame = tk.Frame(frame, bg=row_bg)
                btn_frame.pack(side="top", fill="x", padx=2)
                
                tk.Button(btn_frame, text="OPEN", command=lambda p=path: self._open_file_and_monitor(p), width=6, height=1, font=("Helvetica", 8), bg="#c8e6c9").pack(side="left", padx=(0,2))
                tk.Button(btn_frame, text="DIG", command=lambda p=path: self.dig_document(p), width=6, height=1, font=("Helvetica", 8), bg="#bbdefb").pack(side="left", padx=(2,2))
                tk.Button(btn_frame, text="FOLDER", command=lambda p=path: open_file(os.path.dirname(p)), width=6, height=1, font=("Helvetica", 8), bg="#ffe0b2").pack(side="left", padx=(2,2))
                
                # FIND button uses find_in_tree
                tk.Button(btn_frame, text="FIND", command=lambda p=path: self.find_in_tree(p), width=7, height=1, font=("Helvetica", 8), bg="#e1bee7").pack(side="left", padx=(2,2))
                
                # Apply the fg=txt_color so the score label is readable
                tk.Label(btn_frame, text=f"{score:.4f}", font=("Helvetica", 10), bg=row_bg, fg=txt_color, anchor="w").pack(side="left", padx=(2,0))

                # Pack MOVE button next to the score
                tk.Button(btn_frame, text="MOVE", command=lambda p=path: self.move_selected_file(p), width=6, height=1, font=("Helvetica", 8), bg="#ffcdd2").pack(side="left", padx=(2,0))

                # Pack the full relative folder path on the absolute right (excluding root)
                if parent_dir:
                    try:
                        rel_path = Path(path).parent.relative_to(parent_dir)
                        if str(rel_path) == ".":
                            display_path = ""  # It's directly in the root directory
                        else:
                            display_path = str(rel_path)
                    except ValueError:
                        display_path = os.path.basename(os.path.dirname(path))
                else:
                    display_path = os.path.basename(os.path.dirname(path))
                    
                # Truncate from the left if it's too long, so the final folder remains visible
                max_path_len = 40
                if len(display_path) > max_path_len:
                    display_path = "..." + display_path[-(max_path_len-3):]
                
                # Only pack the label if there's actually a subfolder to show
                if display_path:
                    tk.Label(btn_frame, text=f"({display_path})", font=("Helvetica", 8, "italic"), bg=row_bg, fg=txt_color, anchor="e").pack(side="right", padx=(2, 5))

                # Filename label reverted to a single line
                label = tk.Label(frame, text=os.path.basename(path), font=self._rec_font_normal, bg=row_bg, fg=txt_color, anchor="w")
                if self.current_active_file == path:
                    label.config(font=self._rec_font_bold)
                    self.current_active_label = label
                label.pack(fill="x", side="top", padx=2, pady=(0,2))

                label.bind("<Button-1>", lambda event, p=path: self.update_about_document(event, p))
                label.bind("<Double-1>", lambda event, p=path: self.find_in_tree(p))

    def _on_erase_sidebar_directory(self, dir_path):
        self.erase_directory(dir_path)

    def move_selected_file(self, target_path):
        src = self.current_document
        if not src or not os.path.isfile(src):
            messagebox.showerror("Error", "No valid file selected to move.")
            return
        dest_dir = os.path.dirname(target_path)
        if not os.path.isdir(dest_dir):
            try:
                os.makedirs(dest_dir, exist_ok=True)
            except Exception as e:
                messagebox.showerror("Error", f"Cannot create destination directory: {e}")
                return
        dest_path = os.path.join(dest_dir, os.path.basename(src))
        try:
            shutil.move(src, dest_path)
        except shutil.Error as e:
            messagebox.showerror("Move Error", f"Failed to move file (shutil error): {e}")
            return
        except Exception as e:
            messagebox.showerror("Move Error", f"Unexpected error moving file: {e}")
            return
        self.output_text.config(state="normal")
        self.output_text.delete("1.0", tk.END)
        self.output_text.insert(tk.END, f"Document moved to {dest_path}")
        self.output_text.config(state="disabled")

    def open_tree_to_path(self, target_path):
        for iid in self.dir_tree.get_children(""):
            self.dir_tree.item(iid, open=False)

        candidates = [d for d in self.directories if target_path.startswith(d)]
        if not candidates:
            return
        root_dir = max(candidates, key=len)

        root_iid = None
        for iid in self.dir_tree.get_children(""):
            if self.dir_tree.item(iid, "values")[0] == root_dir:
                root_iid = iid
                break
        if root_iid is None:
            return

        rel = os.path.relpath(target_path, root_dir)
        parts = rel.split(os.sep)
        folder_segments, filename = parts[:-1], parts[-1]

        current = root_iid
        for seg in folder_segments:
            self._expand_and_populate(current)
            for child in self.dir_tree.get_children(current):
                if self.dir_tree.item(child, "text") == seg:
                    current = child
                    break
            else:
                return

        self._expand_and_populate(current)
        for child in self.dir_tree.get_children(current):
            if self.dir_tree.item(child, "text") == filename:
                self.dir_tree.see(child)
                self.dir_tree.selection_set(child)
                self.dir_tree.focus(child)
                break

    def recommend_similar_files(self, current_path, current_text=None):
        import threading
        
        print(f"[LOG] Analyzing: {current_path}")
       
        self.clear_shortcut_buttons()
        if hasattr(self, 'status_label') and self.status_label.winfo_exists():
            self.status_label.config(text=f"Generating recommendations...")
            self.root.update_idletasks()

        if current_text is None:
            current_text = cache_generator.extract_text(current_path)

        def recommendation_task():
            try:
                # CALL 1: Fetch the Amplified Results
                ranked_amplified = recommender.get_recommendations(
                    current_text=current_text,
                    cache=self.cache,
                    directory_active_status=self.directory_active_status,
                    amplifiers=self.amplifiers,
                    silencers=self.silencers,
                    rest_part=getattr(self, "rest_part", ""),
                    word_count_threshold=self.length_penalty.get("threshold", 3000),
                    length_factor=self.length_penalty.get("factor", 0.9)
                )

                # CALL 2: Fetch the Base Results (bypassing multipliers)
                ranked_base = recommender.get_recommendations(
                    current_text=current_text,
                    cache=self.cache,
                    directory_active_status=self.directory_active_status,
                    amplifiers={"words": [], "factor": 1.0},
                    silencers=self.silencers,
                    rest_part="",
                    word_count_threshold=self.length_penalty.get("threshold", 3000),
                    length_factor=self.length_penalty.get("factor", 0.9)
                )

                def finalize():
                    # Store both lists in memory
                    self.recommended_docs_amplified = ranked_amplified
                    self.recommended_docs_base = ranked_base
                    
                    # Tell the UI to draw the correct one
                    self.refresh_recommendation_buttons()
                    
                    if hasattr(self, 'status_label') and self.status_label.winfo_exists():
                        self.status_label.config(text="Recommendations ready.")

                self.root.after(0, finalize)
                
            except Exception as e:
                print(f"[ERROR] Recommendation task failed: {e}")
                self.root.after(0, lambda: self.status_label.config(text="Recommendation failed."))

        threading.Thread(target=recommendation_task, daemon=True).start()
    
    # ----------------------------------------------------------------------
    # YEAR LOOKBACK FEATURE WITH CALENDAR HEATMAP (PDF)
    # ----------------------------------------------------------------------
    def generate_year_lookback(self):
        """Generates a PDF summary including a daily heatmap and opens it."""
        # 1. Check dependencies
        try:
            import fitz  # PyMuPDF
        except ImportError:
            messagebox.showerror("Error", "PyMuPDF is missing.\nPlease run: pip install pymupdf")
            return

        # 2. Run generation with error catching
        try:
            save_dir = str(Path(__file__).parent.resolve())
            
            # Show a loading status (optional but helpful)
            self.status_label.config(text="Generating Year Lookback PDF...")
            self.root.update_idletasks()

            pdf_path = year_lookback.generate_report(self.cache, self.directory_active_status, save_dir)
            
            if pdf_path and os.path.exists(pdf_path):
                self.status_label.config(text=f"Generated: {os.path.basename(pdf_path)}")
                open_file(pdf_path)
            else:
                self.status_label.config(text="Generation skipped (no data).")
                messagebox.showinfo("Lookback", "No files found for the current year to generate a report.")
                
        except PermissionError:
            messagebox.showerror("Permission Error", f"Could not write PDF to:\n{save_dir}\n\nTry moving the app to a folder like Documents.")
        except Exception as e:
            # Catch unexpected crashes and show them
            messagebox.showerror("Generation Error", f"An error occurred:\n{e}")
            print(f"[ERROR] Year Lookback failed: {e}")

    def refresh_toggle_buttons(self):
        for frm in self.toggle_frames.values():
            frm.destroy()
        self.toggle_buttons.clear()
        self.toggle_frames.clear()

        self.toggle_btns_frame.columnconfigure(0, weight=1)
        self.toggle_btns_frame.columnconfigure(1, weight=1)

        sorted_dirs = sorted(self.directories)
        for i, directory in enumerate(sorted_dirs):
            is_active = self.directory_active_status.get(directory, True)
            base_color = self.directory_colors.get(directory, "#ffffff")

            display_name = os.path.basename(directory)
            if not display_name and Path(directory).is_dir() and Path(directory).drive:
                display_name = directory
            elif not display_name:
                display_name = directory

            row = i // 2
            column = i % 2
            
            frm = ttk.Frame(self.toggle_btns_frame)
            frm.grid(row=row, column=column, sticky="ew", padx=2, pady=1)

            # Assign assigned color if ON, gray if OFF. Calculate optimal text color.
            btn_color = base_color if is_active else "#f0f0f0"
            txt_color = self._get_contrast_color(btn_color) if is_active else "black"
            
            btn = tk.Button(frm,
                text=display_name,
                bg=btn_color,
                fg=txt_color,
                relief="raised",
                command=lambda d=directory: self._on_sidebar_toggle(d)
            )
            btn.pack(side="left", fill="x", expand=True)

            btn.bind("<Button-3>", lambda e, d=directory: self._change_dir_color(d))

            erase = tk.Button(frm,
                text="X",
                bg="#ff4444", fg="white", relief="raised",
                width=2,
                command=lambda d=directory: self._on_erase_sidebar_directory(d)
            )
            erase.pack(side="right", padx=(2,0))

            self.toggle_frames[directory]  = frm
            self.toggle_buttons[directory] = btn

    def toggle_directory(self, directory):
        self._clear_active_preset()
        current = self.directory_active_status.get(directory, True)
        new_status = not current
        self.directory_active_status[directory] = new_status
        self.save_dashboard()
        
        # Color goes to default system gray when off
        new_color = self.directory_colors.get(directory, "#ffffff") if new_status else "#f0f0f0"
        txt_color = self._get_contrast_color(new_color) if new_status else "black"

        btn = self.toggle_buttons.get(directory)
        if not btn:
            try:
                target_res = str(Path(directory).resolve())
                for k, b in self.toggle_buttons.items():
                    if str(Path(k).resolve()) == target_res:
                        btn = b
                        break
            except Exception:
                pass
        
        if btn:
            btn.config(bg=new_color, fg=txt_color)
        else:
            self.refresh_toggle_buttons()

# ---------- Public entrypoint for modular startup ----------
def launch_bora(root=None, maximize=True):
    """
    Create the Tk root (if needed), initialize the UI, and optionally maximize.
    Returns the DocumentRecommenderApp instance.
    """
    hide_console_window()
    created_root = False
    if root is None:
        root = tk.Tk()
        created_root = True

    set_app_icon(root)  # Set the window icon

    app = DocumentRecommenderApp(root)
    if maximize:
        maximize_window(root)

    if created_root:
        app.load_cache() # Safely load cache if launched directly
        root.mainloop()
        
    return app

# ---------- CLI entry: run the app from bora.py ----------
if __name__ == "__main__":
    import time
    
    update_splash("Checking environment...", 10)

    # Check Dependencies
    if not getattr(sys, 'frozen', False):
        try:
            import deps
            deps.ensure_dependencies(status_callback=lambda msg: update_splash(msg, progress_var.get() + 5))
        except Exception as e:
            print(f"[DEPS] bootstrap skipped or failed: {e}")
    else:
        update_splash("Running as executable...", 30)

    update_splash("Building user interface...", 60)

    # Launch App UI (Notice maximize=False to prevent the frozen window from popping up early)
    app = launch_bora(root=root, maximize=False)

    # Load Cache (This is the heavy step!)
    update_splash("Loading cached files...", 85)
    app.load_cache()

    update_splash("Ready.", 100)
    time.sleep(0.3) # Let the user see 100% completion before revealing the app

    # Reveal
    try:
        splash.destroy()
    except Exception:
        pass
        
    maximize_window(root) # <--- Maximize safely here, now that everything is loaded!
    root.deiconify()
    root.lift()
    root.focus_force()
    root.mainloop()