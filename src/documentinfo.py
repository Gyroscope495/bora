import os
import re
import io
import fitz  # PyMuPDF
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from html import unescape
import tkinter as tk
import tkinter.font as tkfont
from datetime import datetime
from pathlib import Path
import urllib.parse
import subprocess
from PIL import Image, ImageTk, ImageOps
import date_extraction

# --- Constants and Helper Functions for Document Info Display ---

HIGHLIGHT_TO_HEX = {
    WD_COLOR_INDEX.YELLOW      : "#FFFF00",
    WD_COLOR_INDEX.BRIGHT_GREEN: "#00FF00",
    WD_COLOR_INDEX.TURQUOISE   : "#40E0D0",
    WD_COLOR_INDEX.PINK        : "#FFC0CB",
    WD_COLOR_INDEX.BLUE        : "#0000FF",
    WD_COLOR_INDEX.RED         : "#FF0000",
    WD_COLOR_INDEX.DARK_BLUE   : "#00008B",
    WD_COLOR_INDEX.TEAL        : "#008080",
    WD_COLOR_INDEX.GREEN       : "#008000",
    WD_COLOR_INDEX.VIOLET      : "#EE82EE",
    WD_COLOR_INDEX.DARK_RED    : "#8B0000",
    WD_COLOR_INDEX.DARK_YELLOW : "#CCCC00",
    WD_COLOR_INDEX.GRAY_50     : "#808080",
    WD_COLOR_INDEX.GRAY_25     : "#C0C0C0",
    WD_COLOR_INDEX.BLACK       : "#000000",
    WD_COLOR_INDEX.WHITE       : "#FFFFFF",
}

def _compute_age(publish_date):
    now = datetime.now()
    years = now.year - publish_date.year
    months = now.month - publish_date.month
    days = now.day - publish_date.day
    if days < 0:
        months -= 1
        prev_month = (now.month - 1) or 12
        prev_year = now.year if now.month > 1 else now.year - 1
        days_in_prev = (datetime(now.year, now.month, 1) - datetime(prev_year, prev_month, 1)).days
        days += days_in_prev
    if months < 0:
        years -= 1
        months += 12
    return years, months, days

def _parse_css_style(style_str):
    style_dict = {}
    if style_str:
        properties = style_str.split(";")
        for prop in properties:
            if ":" in prop:
                key, value = prop.split(":", 1)
                style_dict[key.strip().lower()] = value.strip()
    return style_dict

def _calculate_best_closeness(text_content, search_words):
    if not search_words or len(search_words) < 2:
        return None

    words = re.findall(r'\b\w+\b', text_content.lower())
    search_words_set = set(w.lower() for w in search_words)
    
    text_words_set = set(words)
    if not search_words_set.issubset(text_words_set):
        return None

    occurrences = []
    for i, word in enumerate(words):
        if word in search_words_set:
            occurrences.append({'idx': i, 'word': word})

    if not occurrences:
        return None

    min_score = float('inf')
    window = [] 
    word_counts_in_window = {}
    
    for right_item in occurrences:
        window.append(right_item)
        word_counts_in_window[right_item['word']] = word_counts_in_window.get(right_item['word'], 0) + 1

        while len(word_counts_in_window) == len(search_words_set):
            start_idx = window[0]['idx']
            end_idx = window[-1]['idx']
            
            span_length = end_idx - start_idx + 1
            score = span_length - len(search_words_set)
            if score < min_score:
                min_score = score

            left_item = window.pop(0)
            word_counts_in_window[left_item['word']] -= 1
            if word_counts_in_window[left_item['word']] == 0:
                del word_counts_in_window[left_item['word']]
                
    return min_score if min_score != float('inf') else None


def retrieve_html_highlights(path):
    highlights = []
    try:
        content = Path(path).read_text(errors="ignore")
        soup = BeautifulSoup(content, "html.parser")
        spans = soup.find_all("span", style=lambda s: s and "background-color:" in s.lower())
        for span in spans:
            style = span.get("style", "")
            style_dict = _parse_css_style(style)
            try:
                color_part = style.lower().split("background-color:")[1]
                color_hex = color_part.split(";")[0].strip()
            except Exception:
                color_hex = "#ffff99"
            text = unescape(span.get_text(strip=True))
            if text:
                highlights.append((color_hex, text, style_dict, None)) 
        return highlights
    except Exception as e:
        print(f"[ERROR] retrieve_html_highlights: {e}")
        return []

def retrieve_docx_highlights(path):
    highlights = []
    try:
        doc = docx.Document(path)

        def _remove_hyperlinks_in_element(element):
            for hyperlink in element.findall(qn('w:hyperlink')):
                parent = hyperlink.getparent()
                if parent is not None:
                    runs = hyperlink.findall(qn('w:r'))
                    for run_element in runs:
                        rPr = run_element.find(qn('w:rPr'))
                        if rPr is not None:
                            rStyle = rPr.find(qn('w:rStyle'))
                            if rStyle is not None and rStyle.get(qn('w:val')) == 'Hyperlink':
                                rPr.remove(rStyle)
                        parent.insert(parent.index(hyperlink), run_element)
                    parent.remove(hyperlink)

        for para in doc.paragraphs:
            _remove_hyperlinks_in_element(para._p)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _remove_hyperlinks_in_element(para._p)

        def _norm_space(s: str) -> str:
            return " ".join(s.split())

        def _flush_current(merged_color, merged_text_parts):
            if merged_color is not None and merged_text_parts:
                hex_color = HIGHLIGHT_TO_HEX.get(merged_color, "#FFFF00")
                style_dict = {
                    "highlight_index": merged_color,
                    "highlight_name": getattr(merged_color, "name", str(merged_color)),
                }
                text_out = " ".join(p for p in merged_text_parts if p)
                if text_out.strip():
                    highlights.append((hex_color, text_out.strip(), style_dict, None))

        def process_runs(runs):
            merged_color = None
            merged_text_parts = []
            for run in runs:
                text_norm = _norm_space(run.text or "")
                if not text_norm:
                    continue
                color_enum = run.font.highlight_color
                if color_enum is not None:
                    if color_enum == merged_color:
                        merged_text_parts.append(text_norm)
                    else:
                        _flush_current(merged_color, merged_text_parts)
                        merged_color = color_enum
                        merged_text_parts = [text_norm]
                else: 
                    _flush_current(merged_color, merged_text_parts)
                    merged_color = None
                    merged_text_parts = []
            _flush_current(merged_color, merged_text_parts)

        for para in doc.paragraphs:
            process_runs(para.runs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        process_runs(p.runs)

        return highlights
    except Exception as e:
        print(f"[ERROR] retrieve_docx_highlights: {e}")
        return []

def retrieve_pdf_highlights(path):
    highlights = []
    try:
        doc = fitz.open(path)
        for page in doc:
            for annot in page.annots(types=[fitz.PDF_ANNOT_HIGHLIGHT]):
                quadpoints = annot.vertices
                text = ""
                if quadpoints:
                    words = page.get_text("words")
                    for i in range(0, len(quadpoints), 4):
                        rect = fitz.Quad(quadpoints[i:i+4]).rect
                        for w in words:
                            word_rect = fitz.Rect(w[:4])
                            if rect.intersects(word_rect):
                                text += w[4] + " "
                    text = " ".join(text.split()).strip()

                color = annot.colors["stroke"] if annot.colors else (1, 1, 0)
                hex_color = '#%02x%02x%02x' % tuple(int(c * 255) for c in color)

                if text:
                    style_dict = {"annot_type": "highlight", "color_rgb": color}
                    highlights.append((hex_color, text, style_dict, page.number + 1))
        return highlights
    except Exception as e:
        print(f"[ERROR] retrieve_pdf_highlights: {e}")
        return []

# --- Image Retrieval & Metadata Functions ---

def retrieve_pdf_images(path, max_images=5):
    images = []
    try:
        doc = fitz.open(path)
        for page_num in range(len(doc)):
            page = doc[page_num]
            for img in page.get_images(full=True):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                images.append((image_bytes, page_num + 1))
                if len(images) >= max_images:
                    return images
    except Exception as e:
        print(f"[ERROR] retrieve_pdf_images: {e}")
    return images

def retrieve_docx_images(path, max_images=5):
    images = []
    try:
        doc = docx.Document(path)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                images.append((rel.target_part.blob, None))
                if len(images) >= max_images:
                    return images
    except Exception as e:
        print(f"[ERROR] retrieve_docx_images: {e}")
    return images

def retrieve_image_metadata(path):
    """Extracts transcription/metadata embedded in JPEG/PNG files."""
    extracted_text = []
    try:
        with Image.open(path) as img:
            # 1. Check for PNG textual chunks
            if 'Description' in img.info:
                extracted_text.append(str(img.info['Description']))
            if 'Title' in img.info:
                extracted_text.append(str(img.info['Title']))

            # 2. Check EXIF data for JPEGs
            try:
                exif = img.getexif()
                if exif:
                    # Tag 270: ImageDescription
                    if 270 in exif:
                        val = exif[270]
                        if isinstance(val, bytes):
                            val = val.decode('utf-8', errors='ignore')
                        extracted_text.append(str(val))
                    
                    # Tag 40091: XPTitle (Windows Properties Title)
                    if 40091 in exif:
                        val = exif[40091]
                        if isinstance(val, bytes):
                            # Windows XP properties are generally UTF-16LE null-terminated
                            val = val.decode('utf-16le', errors='ignore').rstrip('\x00')
                        extracted_text.append(str(val))
            except Exception as e:
                pass # EXIF data might not exist or be readable
                
    except Exception as e:
        print(f"[ERROR] retrieve_image_metadata: {e}")
        
    # Deduplicate in case multiple tags hold the same transcription
    unique_text = list(dict.fromkeys(extracted_text))
    return "\n\n".join(unique_text).strip()

def show_full_image(img_bytes, title="Image Viewer"):
    """Spawns a new Tkinter Toplevel window to display the full resolution image."""
    top = tk.Toplevel()
    top.title(title)
    
    try:
        pil_img = Image.open(io.BytesIO(img_bytes))
        
        # Apply EXIF rotation to correct orientation
        pil_img = ImageOps.exif_transpose(pil_img)
        
        # Scale down if it exceeds the screen dimensions
        screen_width = top.winfo_screenwidth() - 100
        screen_height = top.winfo_screenheight() - 100
        
        w, h = pil_img.size
        if w > screen_width or h > screen_height:
            ratio = min(screen_width / w, screen_height / h)
            new_w = int(w * ratio)
            new_h = int(h * ratio)
            pil_img = pil_img.resize((new_w, new_h), Image.Resampling.LANCZOS)
            
        tk_img = ImageTk.PhotoImage(pil_img)
        
        lbl = tk.Label(top, image=tk_img, bg="gray")
        lbl.image = tk_img  # Store reference
        lbl.pack(fill=tk.BOTH, expand=True)
        
        lbl.bind("<Button-1>", lambda e: top.destroy())
        top.bind("<Escape>", lambda e: top.destroy())
        top.focus_force()
        
    except Exception as e:
        print(f"Failed to open full image: {e}")
        top.destroy()

# --- New Helper Functions ---

def clear_doc_info(output_text_widget, published_var):
    output_text_widget.config(state="normal")
    output_text_widget.delete("1.0", tk.END)
    output_text_widget.config(state="disabled")
    if published_var:
        published_var.set("Published: —")

def display_summary_text(output_text_widget, summary_text):
    output_text_widget.config(state="normal")
    output_text_widget.delete("1.0", tk.END)
    
    base_font = tkfont.Font(font=output_text_widget.cget("font"))

    if "mindmap_header" not in output_text_widget.tag_names():
        header_font = base_font.copy()
        header_font.configure(weight="bold", underline=True, size=12)
        output_text_widget.tag_configure("mindmap_header", font=header_font, justify='center')
    
    if "mindmap_folder" not in output_text_widget.tag_names():
        folder_font = base_font.copy()
        folder_font.configure(weight="bold")
        output_text_widget.tag_configure("mindmap_folder", font=folder_font)

    lines = summary_text.split('\n')
    
    if lines:
        output_text_widget.insert(tk.END, lines[0] + '\n\n', "mindmap_header")
        for line in lines[1:]:
            if line.strip():
                if line.strip().startswith("📁"):
                    output_text_widget.insert(tk.END, line + '\n', "mindmap_folder")
                else:
                    output_text_widget.insert(tk.END, line + '\n')
            else:
                output_text_widget.insert(tk.END, '\n')

    output_text_widget.config(state="disabled")

def _get_text_color_for_bg(hex_color):
    hex_color = hex_color.lstrip('#')
    if len(hex_color) != 6:
        return "#000000"
    try:
        r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        luminance = (0.299 * r + 0.587 * g + 0.114 * b) / 255
        return "#000000" if luminance > 0.5 else "#FFFFFF"
    except ValueError:
        return "#000000"
    
def open_pdf_in_edge(filepath, search_text, page=None):
    clean_text = search_text.strip()
    encoded_text = urllib.parse.quote(clean_text)
    
    abs_path = os.path.abspath(filepath).replace('\\', '/')
    file_url = f"file:///{abs_path}"
    
    fragments = []
    if page:
        fragments.append(f"page={page}")
    if encoded_text:
        fragments.append(f"search={encoded_text}")
        
    if fragments:
        file_url += "#" + "&".join(fragments)
        
    try:
        command = f'start "" msedge "{file_url}"'
        subprocess.Popen(command, shell=True)
    except Exception as e:
        print(f"Failed to open Edge: {e}")

# --- Main public function ---

def display_doc_info(output_text_widget, published_var, file_path, file_text, search_query, font_size):
    MAX_MATCHES = 20
    is_pdf = file_path.lower().endswith(".pdf")
    ext = os.path.splitext(file_path)[1].lower()
    
    IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.tiff'}

    age_str = ""
    pub_date = date_extraction.get_datetime_from_path(file_path)
    if pub_date:
        try:
            y, m, d = _compute_age(pub_date)
            age_str = f"{y} years, {m} months, {d} days ago"
        except Exception:
            pass

    phrases, terms = [], []
    if search_query:
        for match in re.finditer(r'"([^"]+)"|(\S+)', search_query):
            if match.group(1):
                phrases.append(match.group(1))
            else:
                term = match.group(2)
                if term.lower() != 'ayo':
                    terms.append(term)
    tokens = phrases + terms

    def _whole_word_re(tok: str) -> str:
        return r'\b' + re.escape(tok) + r'\b'

    file_text = file_text or ""
    sentences = re.split(r'(?<=[\.\!?])\s+', file_text)
    unique_matches = []

    if len(tokens) > 1:
        scored_sentences = []
        for sent in sentences:
            clean_sent = " ".join(sent.split())
            if not clean_sent: 
                continue
                
            found_tokens_in_sent = {t for t in tokens if re.search(_whole_word_re(t), clean_sent, re.IGNORECASE)}
            if not found_tokens_in_sent: continue

            num_unique = len(found_tokens_in_sent)
            closeness_score = _calculate_best_closeness(clean_sent, tokens)
            
            final_closeness = closeness_score if closeness_score is not None else float('inf')
            score = (-num_unique, final_closeness)
            scored_sentences.append((clean_sent, score))
        
        scored_sentences.sort(key=lambda x: x[1])
        
        seen = set()
        for sent, score in scored_sentences:
            if sent not in seen:
                unique_matches.append(sent)
                seen.add(sent)
                
    elif tokens:
        for token in tokens:
            for sent in sentences:
                clean_sent = " ".join(sent.split())
                if not clean_sent:
                    continue
                    
                if re.search(_whole_word_re(token), clean_sent, re.IGNORECASE) and clean_sent not in unique_matches:
                    unique_matches.append(clean_sent)

    unique_matches = unique_matches[:MAX_MATCHES]

    highlights = []
    images = []
    image_transcription = ""
    
    if ext == ".html":
        highlights = retrieve_html_highlights(file_path)
    elif ext == ".docx":
        highlights = retrieve_docx_highlights(file_path)
        images = retrieve_docx_images(file_path, max_images=5)
    elif ext == ".pdf":
        highlights = retrieve_pdf_highlights(file_path)
        images = retrieve_pdf_images(file_path, max_images=5)
    elif ext in IMAGE_EXTENSIONS:
        try:
            with open(file_path, "rb") as img_file:
                images.append((img_file.read(), None))
            # Extract metadata from the standalone image
            image_transcription = retrieve_image_metadata(file_path)
        except Exception as e:
            print(f"Failed to load standalone image {file_path}: {e}")

    if age_str:
        published_var.set("Published: " + age_str.strip())
    else:
        published_var.set("Published: —")

    output_text_widget.config(state="normal")
    output_text_widget.delete("1.0", tk.END)
    
    # Establish a persistent list for our images to prevent GC
    if not hasattr(output_text_widget, 'image_refs'):
        output_text_widget.image_refs = []
    output_text_widget.image_refs.clear()

    # --- Setup Fonts ---
    base_font = tkfont.Font(font=output_text_widget.cget("font"))
    base_font.configure(size=font_size)
    
    header_font = base_font.copy()
    header_font.configure(size=font_size + 2, weight="bold")

    bold_font = base_font.copy()
    bold_font.configure(weight="bold")
    
    italic_font = base_font.copy()
    italic_font.configure(slant="italic")

    # --- Layout Tags ---
    output_text_widget.tag_configure("section_header", font=header_font, foreground="#2C3E50", spacing1=15, spacing3=10)
    output_text_widget.tag_configure("match_odd", background="#FFFFFF", lmargin1=20, lmargin2=20, spacing1=5, spacing3=5)
    output_text_widget.tag_configure("match_even", background="#F4F6F7", lmargin1=20, lmargin2=20, spacing1=5, spacing3=5)
    output_text_widget.tag_configure("excerpt_text", foreground="#4A4A4A", lmargin1=15, lmargin2=15, spacing1=5, spacing3=5)
    output_text_widget.tag_configure("image_spacing", lmargin1=20, lmargin2=20, spacing1=10, spacing3=10)
    output_text_widget.tag_configure("click_hint", font=italic_font, foreground="#0066CC", lmargin1=15, lmargin2=15)
    output_text_widget.tag_configure("hyperlink_style", foreground="#0066CC", underline=True)

    def insert_hyperlink(widget, link_text, action_func, parent_tags):
        tag_name = f"link_{id(action_func)}_{hash(link_text)}"
        combined_tags = parent_tags + ("hyperlink_style", tag_name)
        widget.insert(tk.END, link_text, combined_tags)
        widget.tag_bind(tag_name, "<Button-1>", lambda e, f=action_func: f())
        widget.tag_bind(tag_name, "<Enter>", lambda e: widget.config(cursor="hand2"))
        widget.tag_bind(tag_name, "<Leave>", lambda e: widget.config(cursor=""))

    pdf_doc = None
    if is_pdf:
        try:
            pdf_doc = fitz.open(file_path)
        except Exception:
            pass

    def _find_match_page(sentence):
        if not pdf_doc: return None
        target = re.sub(r'\s+', ' ', sentence).strip()
        if not target: return None
        target_sub = target[:40].lower() 
        for page in pdf_doc:
            page_text = re.sub(r'\s+', ' ', page.get_text("text")).lower()
            if target_sub in page_text:
                return page.number + 1
        return None

    if unique_matches:
        output_text_widget.insert(tk.END, "🔍 Search Matches\n", "section_header")
        for idx, sent in enumerate(unique_matches, 1):
            tag = "match_even" if idx % 2 == 0 else "match_odd"
            page_num = None
            if is_pdf:
                page_num = _find_match_page(sent)
            
            prefix = f"Page {page_num}: " if page_num is not None else "• "
            output_text_widget.insert(tk.END, f"{prefix}{sent}", tag)
            
            if is_pdf:
                output_text_widget.insert(tk.END, "  ", tag)
                best_token = tokens[0] if tokens else ""
                for t in tokens:
                    if re.search(_whole_word_re(t), sent, re.IGNORECASE):
                        best_token = t
                        break
                edge_search_word = ""
                if best_token:
                    word_match = re.search(r'\w+', best_token)
                    if word_match:
                        edge_search_word = word_match.group(0)
                        
                insert_hyperlink(output_text_widget, "[Find in document]", 
                                 lambda w=edge_search_word, p=page_num: open_pdf_in_edge(file_path, w, p), 
                                 (tag,))
            output_text_widget.insert(tk.END, "\n", tag)

    if pdf_doc:
        pdf_doc.close()

    if ext not in IMAGE_EXTENSIONS:
        output_text_widget.insert(tk.END, "🖍️ Highlights\n", "section_header")
        if highlights:
            for idx, (color_hex, hl_text, style_dict, page_num) in enumerate(highlights, 1):
                text_color = _get_text_color_for_bg(color_hex)
                tag_name = f"hl_{color_hex}_{abs(hash(hl_text))}"
                if tag_name not in output_text_widget.tag_names():
                    tag_font = base_font.copy()
                    if style_dict.get("bold"): tag_font.configure(weight="bold")
                    if style_dict.get("italic"): tag_font.configure(slant="italic")
                    output_text_widget.tag_configure(tag_name, font=tag_font, background=color_hex, 
                                                     foreground=text_color, lmargin1=20, lmargin2=20, 
                                                     spacing1=8, spacing3=8)
                    
                prefix = f"Page {page_num}: " if page_num is not None else "• "
                output_text_widget.insert(tk.END, f"{prefix}{hl_text}", tag_name)
                
                if is_pdf:
                    output_text_widget.insert(tk.END, "  ", tag_name)
                    insert_hyperlink(output_text_widget, "[Find in document]", lambda t=hl_text, p=page_num: open_pdf_in_edge(file_path, t, p), (tag_name,))
                output_text_widget.insert(tk.END, "\n", tag_name)
        else:
            output_text_widget.insert(tk.END, "No highlights found.\n", "excerpt_text")

    if images:
        header_text = "🖼️ Image Preview\n" if ext in IMAGE_EXTENSIONS else "🖼️ Extracted Images\n"
        output_text_widget.insert(tk.END, header_text, "section_header")
        output_text_widget.insert(tk.END, "(Click an image to view it full size)\n", "click_hint")
        
        max_img_width = 300 
        for img_bytes, page_num in images:
            try:
                pil_img = Image.open(io.BytesIO(img_bytes))
                
                # Apply EXIF rotation to correct orientation
                pil_img = ImageOps.exif_transpose(pil_img)

                w_percent = (max_img_width / float(pil_img.size[0]))
                if w_percent < 1.0: 
                    h_size = int((float(pil_img.size[1]) * float(w_percent)))
                    pil_img = pil_img.resize((max_img_width, h_size), Image.Resampling.LANCZOS)
                
                tk_img = ImageTk.PhotoImage(pil_img)
                
                # Append to our persistent list to save it from the garbage collector
                output_text_widget.image_refs.append(tk_img)

                if page_num:
                    output_text_widget.insert(tk.END, f"Page {page_num}:\n", "excerpt_text")
                
                img_label = tk.Label(output_text_widget, image=tk_img, cursor="hand2", bg="white")
                img_label.image = tk_img  
                
                title_str = f"Image from Page {page_num}" if page_num else ("Image Preview" if ext in IMAGE_EXTENSIONS else "Extracted Document Image")
                img_label.bind("<Button-1>", lambda e, b=img_bytes, t=title_str: show_full_image(b, t))
                
                output_text_widget.window_create(tk.END, window=img_label)
                output_text_widget.insert(tk.END, "\n\n", "image_spacing")
            except Exception as e:
                print(f"Failed to render image: {e}")

    # --- Render Extracted Image Metadata / Transcription ---
    if image_transcription:
        output_text_widget.insert(tk.END, "📝 Embedded Transcription & Metadata\n", "section_header")
        output_text_widget.insert(tk.END, image_transcription + "\n\n", "excerpt_text")

    if file_text.strip() and ext not in IMAGE_EXTENSIONS:
        output_text_widget.insert(tk.END, "📄 Original (first 500 words)\n", "section_header")
        words = file_text.split()
        first_words = " ".join(words[:500])
        output_text_widget.insert(tk.END, first_words + "\n", "excerpt_text")

    # --- Search Highlighting ---
    SEARCH_COLORS = [
        "#FFFF00", "#00FFFF", "#FFBAFF", "#00FF00", "#FFB347", 
        "#FFC0CB", "#87CEFA", "#98FB98", "#FFD700", "#FFA07A"
    ]

    full_text_str = output_text_widget.get("1.0", tk.END)
    
    for i, token in enumerate(tokens):
        bg_color = SEARCH_COLORS[i % len(SEARCH_COLORS)]
        fg_color = _get_text_color_for_bg(bg_color)
        tag_name = f"search_bold_{i}"
        
        output_text_widget.tag_configure(
            tag_name, 
            font=bold_font, 
            foreground=fg_color, 
            background=bg_color
        )
        output_text_widget.tag_raise(tag_name)
        for m in re.finditer(_whole_word_re(token), full_text_str, re.IGNORECASE):
            start_idx = f"1.0+{m.start()}c"
            end_idx = f"1.0+{m.end()}c"
            output_text_widget.tag_add(tag_name, start_idx, end_idx)

    output_text_widget.config(state="disabled")